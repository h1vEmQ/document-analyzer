#!/usr/bin/env python
"""
Тестирование автоматической генерации отчетов после анализа
"""

import os
import sys
import django

# Настройка Django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'wara_project.settings')
django.setup()

from django.test import Client
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from documents.models import Document as DjangoDocument
from analysis.models import Comparison
from reports.models import Report
from reports.services import AutoReportGeneratorService, DOCXReportGeneratorService
from docx import Document as DocxDocument
from io import BytesIO

User = get_user_model()

def create_test_docx_v1():
    """Создание первой версии тестового .docx файла"""
    doc = DocxDocument()
    
    # Добавляем заголовок
    doc.add_heading('Отчет о продажах - Версия 1', 0)
    
    # Добавляем параграф
    doc.add_paragraph('Это первая версия отчета о продажах за квартал.')
    
    # Добавляем раздел
    doc.add_heading('Раздел 1: Основные показатели', level=1)
    doc.add_paragraph('Общий объем продаж составил 100 000 рублей.')
    
    # Добавляем таблицу
    doc.add_heading('Таблица продаж', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Заполняем таблицу
    table.cell(0, 0).text = 'Продукт'
    table.cell(0, 1).text = 'Количество'
    table.cell(0, 2).text = 'Сумма'
    
    table.cell(1, 0).text = 'Товар А'
    table.cell(1, 1).text = '10'
    table.cell(1, 2).text = '50 000'
    
    table.cell(2, 0).text = 'Товар Б'
    table.cell(2, 1).text = '5'
    table.cell(2, 2).text = '50 000'
    
    # Добавляем еще один раздел
    doc.add_heading('Раздел 2: Выводы', level=1)
    doc.add_paragraph('Продажи соответствуют плану.')
    
    # Сохраняем в байты
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def create_test_docx_v2():
    """Создание второй версии тестового .docx файла (с изменениями)"""
    doc = DocxDocument()
    
    # Добавляем заголовок
    doc.add_heading('Отчет о продажах - Версия 2', 0)
    
    # Добавляем параграф (изменен)
    doc.add_paragraph('Это обновленная версия отчета о продажах за квартал.')
    
    # Добавляем раздел
    doc.add_heading('Раздел 1: Основные показатели', level=1)
    doc.add_paragraph('Общий объем продаж составил 120 000 рублей.')  # Изменено
    
    # Добавляем таблицу (изменена)
    doc.add_heading('Таблица продаж', level=2)
    table = doc.add_table(rows=4, cols=3)  # Добавлена строка
    table.style = 'Table Grid'
    
    # Заполняем таблицу
    table.cell(0, 0).text = 'Продукт'
    table.cell(0, 1).text = 'Количество'
    table.cell(0, 2).text = 'Сумма'
    
    table.cell(1, 0).text = 'Товар А'
    table.cell(1, 1).text = '12'  # Изменено
    table.cell(1, 2).text = '60 000'  # Изменено
    
    table.cell(2, 0).text = 'Товар Б'
    table.cell(2, 1).text = '5'
    table.cell(2, 2).text = '50 000'
    
    table.cell(3, 0).text = 'Товар В'  # Добавлено
    table.cell(3, 1).text = '2'  # Добавлено
    table.cell(3, 2).text = '10 000'  # Добавлено
    
    # Добавляем еще один раздел (изменен)
    doc.add_heading('Раздел 2: Выводы', level=1)
    doc.add_paragraph('Продажи превысили план на 20%.')  # Изменено
    
    # Добавляем новый раздел
    doc.add_heading('Раздел 3: Рекомендации', level=1)
    doc.add_paragraph('Рекомендуется увеличить производство товара А.')  # Добавлено
    
    # Сохраняем в байты
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def test_auto_reports():
    """Тестирование автоматической генерации отчетов"""
    print("🧪 Тестирование автоматической генерации отчетов")
    print("=" * 60)
    
    # Получаем или создаем тестового пользователя
    user, created = User.objects.get_or_create(
        username='testuser',
        defaults={
            'email': 'test@example.com',
            'first_name': 'Тест',
            'last_name': 'Пользователь',
            'role': 'user'
        }
    )
    if created:
        user.set_password('testpass123')
        user.save()
        print(f"    ✅ Создан пользователь: {user.username}")
    else:
        print(f"    ✅ Используется существующий пользователь: {user.username}")
    
    # Создаем тестовые .docx файлы
    print("\n📄 Создание тестовых документов...")
    docx_v1_content = create_test_docx_v1()
    docx_v2_content = create_test_docx_v2()
    
    test_file_v1 = SimpleUploadedFile(
        "test_report_v1.docx",
        docx_v1_content,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    test_file_v2 = SimpleUploadedFile(
        "test_report_v2.docx",
        docx_v2_content,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    print(f"    ✅ Создан файл v1: {test_file_v1.name} ({len(docx_v1_content)} байт)")
    print(f"    ✅ Создан файл v2: {test_file_v2.name} ({len(docx_v2_content)} байт)")
    
    # Создаем клиент для тестирования
    client = Client()
    client.force_login(user)
    
    # Тест 1: Загрузка документов
    print("\n🔍 Тест 1: Загрузка документов")
    
    # Загружаем первый документ
    response1 = client.post('/upload/', {
        'title': 'Отчет о продажах v1',
        'file': test_file_v1
    })
    
    if response1.status_code == 302:
        print("    ✅ Первый документ загружен")
    else:
        print(f"    ❌ Ошибка загрузки первого документа: {response1.status_code}")
        return
    
    # Загружаем второй документ
    response2 = client.post('/upload/', {
        'title': 'Отчет о продажах v2',
        'file': test_file_v2
    })
    
    if response2.status_code == 302:
        print("    ✅ Второй документ загружен")
    else:
        print(f"    ❌ Ошибка загрузки второго документа: {response2.status_code}")
        return
    
    # Получаем загруженные документы
    documents = DjangoDocument.objects.filter(user=user).order_by('-upload_date')[:2]
    if len(documents) < 2:
        print("    ❌ Недостаточно документов для сравнения")
        return
    
    doc1, doc2 = documents[0], documents[1]
    print(f"    📝 Документ 1: {doc1.title} (ID: {doc1.id})")
    print(f"    📝 Документ 2: {doc2.title} (ID: {doc2.id})")
    
    # Тест 2: Создание сравнения через форму (с автоматической генерацией отчетов)
    print("\n🔍 Тест 2: Создание сравнения с автоматической генерацией отчетов")
    
    response = client.post('/analysis/create/', {
        'title': 'Автоматические отчеты тест',
        'base_document': doc1.id,
        'compared_document': doc2.id
    })
    
    if response.status_code == 302:  # Редирект после успешного создания
        print("    ✅ Сравнение успешно создано")
        
        # Получаем последнее созданное сравнение
        latest_comparison = Comparison.objects.filter(user=user).order_by('-created_date').first()
        if latest_comparison:
            print(f"    📝 Сравнение ID: {latest_comparison.id}")
            print(f"    📝 Название: {latest_comparison.title}")
            print(f"    📝 Статус: {latest_comparison.status}")
            
            # Проверяем, был ли анализ выполнен автоматически
            if latest_comparison.status == 'completed':
                print("    ✅ Анализ автоматически завершен!")
                
                # Проверяем автоматически созданные отчеты
                auto_reports = Report.objects.filter(
                    comparison=latest_comparison,
                    title__contains='Auto-generated'
                ).order_by('generated_date')
                
                print(f"    📊 Автоматически созданных отчетов: {auto_reports.count()}")
                
                for report in auto_reports:
                    print(f"        - {report.title} (ID: {report.id}, Формат: {report.format}, Статус: {report.status})")
                    if report.file:
                        print(f"          Файл: {report.file.name}")
                
                # Проверяем наличие PDF и DOCX отчетов
                pdf_reports = auto_reports.filter(format='pdf')
                docx_reports = auto_reports.filter(format='docx')
                
                if pdf_reports.exists():
                    print("    ✅ PDF отчет автоматически создан")
                else:
                    print("    ❌ PDF отчет не создан")
                
                if docx_reports.exists():
                    print("    ✅ DOCX отчет автоматически создан")
                else:
                    print("    ❌ DOCX отчет не создан")
            else:
                print(f"    ❌ Анализ не был автоматически завершен (статус: {latest_comparison.status})")
        else:
            print("    ❌ Сравнение не найдено в базе данных")
    else:
        print(f"    ❌ Ошибка создания сравнения: статус {response.status_code}")
        if hasattr(response, 'content'):
            print(f"    📄 Содержимое ответа: {response.content.decode()[:200]}...")
    
    # Тест 3: Тестирование сервиса генерации DOCX отчетов
    print("\n🔍 Тест 3: Тестирование сервиса генерации DOCX отчетов")
    
    latest_comparison = Comparison.objects.filter(user=user).order_by('-created_date').first()
    if latest_comparison and latest_comparison.status == 'completed':
        try:
            docx_service = DOCXReportGeneratorService()
            docx_content = docx_service.generate_report(latest_comparison)
            
            print(f"    ✅ DOCX отчет сгенерирован")
            print(f"    📄 Размер DOCX файла: {len(docx_content)} байт")
            
            # Проверяем, что это валидный DOCX файл
            if docx_content.startswith(b'PK'):
                print("    ✅ Файл является валидным DOCX (начинается с PK)")
            else:
                print("    ❌ Файл не является валидным DOCX")
            
        except Exception as e:
            print(f"    ❌ Ошибка генерации DOCX отчета: {e}")
    else:
        print("    ⏭️ Тест DOCX генерации пропущен (нет завершенного сравнения)")
    
    # Тест 4: Тестирование сервиса автоматической генерации отчетов
    print("\n🔍 Тест 4: Тестирование сервиса автоматической генерации отчетов")
    
    if latest_comparison and latest_comparison.status == 'completed':
        try:
            auto_service = AutoReportGeneratorService()
            results = auto_service.generate_auto_reports(latest_comparison)
            
            print(f"    ✅ Сервис автоматической генерации выполнен")
            print(f"    📊 Результаты: {results}")
            
            if results.get('pdf_report'):
                print("    ✅ PDF отчет создан через сервис")
            if results.get('docx_report'):
                print("    ✅ DOCX отчет создан через сервис")
            
            if results.get('errors'):
                print("    ⚠️ Ошибки при генерации:")
                for error in results['errors']:
                    print(f"        - {error}")
            
        except Exception as e:
            print(f"    ❌ Ошибка автоматической генерации отчетов: {e}")
    else:
        print("    ⏭️ Тест автоматической генерации пропущен (нет завершенного сравнения)")
    
    # Тест 5: Проверка интерфейса отчетов
    print("\n🔍 Тест 5: Проверка интерфейса отчетов")
    
    response = client.get('/reports/')
    if response.status_code == 200:
        print("    ✅ Страница отчетов загружена")
        
        content = response.content.decode()
        auto_reports_count = Report.objects.filter(
            user=user,
            title__contains='Auto-generated'
        ).count()
        
        if auto_reports_count > 0:
            print(f"    ✅ Найдено автоматически созданных отчетов: {auto_reports_count}")
            
            # Проверяем отображение отчетов на странице
            if 'Auto-generated' in content:
                print("    ✅ Автоматически созданные отчеты отображаются на странице")
            else:
                print("    ⚠️ Автоматически созданные отчеты не найдены на странице")
        else:
            print("    ❌ Автоматически созданные отчеты не найдены")
    else:
        print(f"    ❌ Ошибка загрузки страницы отчетов: {response.status_code}")
    
    print("\n🎉 Тестирование автоматической генерации отчетов завершено!")
    print("=" * 60)
    
    # Итоговая статистика
    total_reports = Report.objects.filter(user=user).count()
    auto_reports = Report.objects.filter(user=user, title__contains='Auto-generated').count()
    pdf_reports = Report.objects.filter(user=user, format='pdf').count()
    docx_reports = Report.objects.filter(user=user, format='docx').count()
    
    print(f"📊 Итоговая статистика:")
    print(f"    📄 Всего отчетов: {total_reports}")
    print(f"    🤖 Автоматически созданных отчетов: {auto_reports}")
    print(f"    📋 PDF отчетов: {pdf_reports}")
    print(f"    📝 DOCX отчетов: {docx_reports}")
    
    if auto_reports > 0:
        success_rate = (auto_reports / max(total_reports, 1)) * 100
        print(f"    📈 Процент автоматически созданных отчетов: {success_rate:.1f}%")

def main():
    """Основная функция тестирования"""
    try:
        test_auto_reports()
        return True
    except Exception as e:
        print(f"\n❌ Ошибка во время тестирования: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)
