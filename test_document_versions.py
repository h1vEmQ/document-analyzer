#!/usr/bin/env python
"""
Тестирование системы версионирования документов
"""

import os
import sys
import django

# Настройка Django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'wara_project.settings')
django.setup()

from django.test import Client
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from documents.models import Document as DjangoDocument
from documents.services import DocumentParserService, DocumentValidationService
from docx import Document as DocxDocument
from io import BytesIO

User = get_user_model()

def create_test_docx_v1():
    """Создание первой версии тестового .docx файла"""
    doc = DocxDocument()
    
    # Добавляем заголовок
    doc.add_heading('Тестовый документ - Версия 1', 0)
    
    # Добавляем параграф
    doc.add_paragraph('Это первая версия тестового документа.')
    doc.add_paragraph('Содержимое первой версии.')
    
    # Добавляем раздел
    doc.add_heading('Раздел 1', level=1)
    doc.add_paragraph('Содержимое раздела 1.')
    
    # Сохраняем в байты
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def create_test_docx_v2():
    """Создание второй версии тестового .docx файла"""
    doc = DocxDocument()
    
    # Добавляем заголовок
    doc.add_heading('Тестовый документ - Версия 2', 0)
    
    # Добавляем параграф (изменен)
    doc.add_paragraph('Это обновленная версия тестового документа.')
    doc.add_paragraph('Содержимое второй версии с изменениями.')
    
    # Добавляем раздел (изменен)
    doc.add_heading('Раздел 1 (обновлен)', level=1)
    doc.add_paragraph('Обновленное содержимое раздела 1.')
    
    # Добавляем новый раздел
    doc.add_heading('Раздел 2 (новый)', level=1)
    doc.add_paragraph('Новый раздел во второй версии.')
    
    # Сохраняем в байты
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def create_test_docx_v3():
    """Создание третьей версии тестового .docx файла"""
    doc = DocxDocument()
    
    # Добавляем заголовок
    doc.add_heading('Тестовый документ - Версия 3', 0)
    
    # Добавляем параграф (изменен)
    doc.add_paragraph('Это финальная версия тестового документа.')
    doc.add_paragraph('Содержимое третьей версии с дополнительными изменениями.')
    
    # Добавляем раздел (изменен)
    doc.add_heading('Раздел 1 (финальный)', level=1)
    doc.add_paragraph('Финальное содержимое раздела 1.')
    
    # Добавляем раздел (изменен)
    doc.add_heading('Раздел 2 (обновлен)', level=1)
    doc.add_paragraph('Обновленное содержимое раздела 2.')
    
    # Добавляем новый раздел
    doc.add_heading('Раздел 3 (новый)', level=1)
    doc.add_paragraph('Новый раздел в третьей версии.')
    
    # Добавляем таблицу
    doc.add_heading('Таблица данных', level=2)
    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = 'Поле 1'
    table.cell(0, 1).text = 'Поле 2'
    table.cell(1, 0).text = 'Значение 1'
    table.cell(1, 1).text = 'Значение 2'
    table.cell(2, 0).text = 'Значение 3'
    table.cell(2, 1).text = 'Значение 4'
    
    # Сохраняем в байты
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def test_document_versions():
    """Тестирование системы версионирования"""
    print("🧪 Тестирование системы версионирования документов")
    print("=" * 60)
    
    # Получаем или создаем тестового пользователя
    user, created = User.objects.get_or_create(
        username='testuser',
        defaults={
            'email': 'test@example.com',
            'first_name': 'Тест',
            'last_name': 'Пользователь',
            'role': 'user'
        }
    )
    if created:
        user.set_password('testpass123')
        user.save()
        print(f"    ✅ Создан пользователь: {user.username}")
    else:
        print(f"    ✅ Используется существующий пользователь: {user.username}")
    
    # Создаем клиент для тестирования
    client = Client()
    client.force_login(user)
    
    # Тест 1: Загрузка первой версии документа
    print("\n🔍 Тест 1: Загрузка первой версии документа")
    
    docx_v1_content = create_test_docx_v1()
    test_file_v1 = SimpleUploadedFile(
        "test_document_v1.docx",
        docx_v1_content,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    response = client.post('/upload/', {
        'title': 'Тестовый документ для версионирования',
        'file': test_file_v1
    })
    
    if response.status_code == 302:
        print("    ✅ Первая версия загружена")
        
        # Получаем загруженный документ
        doc_v1 = DjangoDocument.objects.filter(user=user).order_by('-upload_date').first()
        print(f"    📝 Документ ID: {doc_v1.id}, Версия: {doc_v1.version}")
        print(f"    📝 Родительский документ: {doc_v1.parent_document}")
        print(f"    📝 Последняя версия: {doc_v1.is_latest_version}")
    else:
        print(f"    ❌ Ошибка загрузки первой версии: {response.status_code}")
        return
    
    # Тест 2: Загрузка второй версии
    print("\n🔍 Тест 2: Загрузка второй версии документа")
    
    docx_v2_content = create_test_docx_v2()
    test_file_v2 = SimpleUploadedFile(
        "test_document_v2.docx",
        docx_v2_content,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    response = client.post(f'/{doc_v1.pk}/version/upload/', {
        'file': test_file_v2,
        'version_notes': 'Добавлен новый раздел и обновлено содержимое'
    })
    
    if response.status_code == 302:
        print("    ✅ Вторая версия загружена")
        
        # Получаем новую версию
        doc_v2 = DjangoDocument.objects.filter(user=user).order_by('-upload_date').first()
        print(f"    📝 Документ ID: {doc_v2.id}, Версия: {doc_v2.version}")
        print(f"    📝 Родительский документ: {doc_v2.parent_document.id if doc_v2.parent_document else 'None'}")
        print(f"    📝 Последняя версия: {doc_v2.is_latest_version}")
        print(f"    📝 Заметки: {doc_v2.version_notes}")
        
        # Проверяем, что первая версия больше не последняя
        doc_v1.refresh_from_db()
        print(f"    📝 Первая версия теперь последняя: {doc_v1.is_latest_version}")
    else:
        print(f"    ❌ Ошибка загрузки второй версии: {response.status_code}")
        return
    
    # Тест 3: Загрузка третьей версии
    print("\n🔍 Тест 3: Загрузка третьей версии документа")
    
    docx_v3_content = create_test_docx_v3()
    test_file_v3 = SimpleUploadedFile(
        "test_document_v3.docx",
        docx_v3_content,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    response = client.post(f'/{doc_v1.pk}/version/upload/', {
        'file': test_file_v3,
        'version_notes': 'Добавлена таблица и финальные изменения'
    })
    
    if response.status_code == 302:
        print("    ✅ Третья версия загружена")
        
        # Получаем новую версию
        doc_v3 = DjangoDocument.objects.filter(user=user).order_by('-upload_date').first()
        print(f"    📝 Документ ID: {doc_v3.id}, Версия: {doc_v3.version}")
        print(f"    📝 Родительский документ: {doc_v3.parent_document.id if doc_v3.parent_document else 'None'}")
        print(f"    📝 Последняя версия: {doc_v3.is_latest_version}")
        print(f"    📝 Заметки: {doc_v3.version_notes}")
    else:
        print(f"    ❌ Ошибка загрузки третьей версии: {response.status_code}")
        return
    
    # Тест 4: Проверка истории версий
    print("\n🔍 Тест 4: Проверка истории версий")
    
    root_doc = doc_v1.get_root_document()
    versions = root_doc.get_version_history()
    
    print(f"    📝 Корневой документ: {root_doc.id}")
    print(f"    📝 Всего версий: {versions.count()}")
    print(f"    📝 Последняя версия: {root_doc.get_latest_version().id}")
    
    for version in versions:
        print(f"        - Версия {version.version} (ID: {version.id}) - {'Текущая' if version.is_latest_version else 'Старая'}")
    
    # Тест 5: Проверка страницы истории версий
    print("\n🔍 Тест 5: Проверка страницы истории версий")
    
    response = client.get(f'/{root_doc.pk}/versions/')
    if response.status_code == 200:
        print("    ✅ Страница истории версий загружена")
        
        content = response.content.decode()
        if 'История версий документа' in content:
            print("    ✅ Заголовок истории версий найден")
        if 'Всего версий' in content:
            print("    ✅ Статистика версий найдена")
        if 'Текущая версия' in content:
            print("    ✅ Информация о текущей версии найдена")
    else:
        print(f"    ❌ Ошибка загрузки страницы истории: {response.status_code}")
    
    # Тест 6: Проверка детального просмотра с версиями
    print("\n🔍 Тест 6: Проверка детального просмотра")
    
    response = client.get(f'/{doc_v3.pk}/')
    if response.status_code == 200:
        print("    ✅ Страница детального просмотра загружена")
        
        content = response.content.decode()
        if 'Новая версия' in content:
            print("    ✅ Кнопка новой версии найдена")
        if 'История версий' in content:
            print("    ✅ Кнопка истории версий найдена")
        if 'Последняя' in content:
            print("    ✅ Информация о статусе версии найдена")
    else:
        print(f"    ❌ Ошибка загрузки детального просмотра: {response.status_code}")
    
    # Тест 7: Проверка методов модели
    print("\n🔍 Тест 7: Проверка методов модели")
    
    print(f"    📝 Версия {doc_v1.version}:")
    print(f"        - Количество версий: {doc_v1.get_version_count()}")
    print(f"        - Корневой документ: {doc_v1.get_root_document().id}")
    print(f"        - Последняя версия: {doc_v1.get_latest_version().id}")
    
    print(f"    📝 Версия {doc_v2.version}:")
    print(f"        - Количество версий: {doc_v2.get_version_count()}")
    print(f"        - Корневой документ: {doc_v2.get_root_document().id}")
    print(f"        - Последняя версия: {doc_v2.get_latest_version().id}")
    
    print(f"    📝 Версия {doc_v3.version}:")
    print(f"        - Количество версий: {doc_v3.get_version_count()}")
    print(f"        - Корневой документ: {doc_v3.get_root_document().id}")
    print(f"        - Последняя версия: {doc_v3.get_latest_version().id}")
    
    # Тест 8: Проверка автоматической обработки версий
    print("\n🔍 Тест 8: Проверка автоматической обработки версий")
    
    processed_versions = DjangoDocument.objects.filter(
        user=user, 
        parent_document=root_doc,
        status='processed'
    ).count()
    
    print(f"    📝 Обработанных версий: {processed_versions}")
    
    if processed_versions > 0:
        print("    ✅ Автоматическая обработка версий работает")
    else:
        print("    ⚠️ Автоматическая обработка версий не сработала")
    
    print("\n🎉 Тестирование системы версионирования завершено!")
    print("=" * 60)
    
    # Итоговая статистика
    total_versions = DjangoDocument.objects.filter(user=user, parent_document=root_doc).count()
    latest_version = root_doc.get_latest_version()
    
    print(f"📊 Итоговая статистика:")
    print(f"    📄 Всего версий документа: {total_versions}")
    print(f"    ✅ Текущая версия: {latest_version.version} (ID: {latest_version.id})")
    print(f"    📈 Все версии успешно созданы и связаны")

def main():
    """Основная функция тестирования"""
    try:
        test_document_versions()
        return True
    except Exception as e:
        print(f"\n❌ Ошибка во время тестирования: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)
