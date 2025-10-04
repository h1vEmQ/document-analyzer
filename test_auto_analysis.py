#!/usr/bin/env python
"""
Тестирование автоматического анализа документов при создании сравнения
"""

import os
import sys
import django

# Настройка Django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'wara_project.settings')
django.setup()

from django.test import Client
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from documents.models import Document as DjangoDocument
from analysis.models import Comparison
from analysis.services import DocumentComparisonService
from docx import Document as DocxDocument
from io import BytesIO

User = get_user_model()

def create_test_docx_v1():
    """Создание первой версии тестового .docx файла"""
    doc = DocxDocument()
    
    # Добавляем заголовок
    doc.add_heading('Отчет о продажах - Версия 1', 0)
    
    # Добавляем параграф
    doc.add_paragraph('Это первая версия отчета о продажах за квартал.')
    
    # Добавляем раздел
    doc.add_heading('Раздел 1: Основные показатели', level=1)
    doc.add_paragraph('Общий объем продаж составил 100 000 рублей.')
    
    # Добавляем таблицу
    doc.add_heading('Таблица продаж', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Заполняем таблицу
    table.cell(0, 0).text = 'Продукт'
    table.cell(0, 1).text = 'Количество'
    table.cell(0, 2).text = 'Сумма'
    
    table.cell(1, 0).text = 'Товар А'
    table.cell(1, 1).text = '10'
    table.cell(1, 2).text = '50 000'
    
    table.cell(2, 0).text = 'Товар Б'
    table.cell(2, 1).text = '5'
    table.cell(2, 2).text = '50 000'
    
    # Добавляем еще один раздел
    doc.add_heading('Раздел 2: Выводы', level=1)
    doc.add_paragraph('Продажи соответствуют плану.')
    
    # Сохраняем в байты
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def create_test_docx_v2():
    """Создание второй версии тестового .docx файла (с изменениями)"""
    doc = DocxDocument()
    
    # Добавляем заголовок
    doc.add_heading('Отчет о продажах - Версия 2', 0)
    
    # Добавляем параграф (изменен)
    doc.add_paragraph('Это обновленная версия отчета о продажах за квартал.')
    
    # Добавляем раздел
    doc.add_heading('Раздел 1: Основные показатели', level=1)
    doc.add_paragraph('Общий объем продаж составил 120 000 рублей.')  # Изменено
    
    # Добавляем таблицу (изменена)
    doc.add_heading('Таблица продаж', level=2)
    table = doc.add_table(rows=4, cols=3)  # Добавлена строка
    table.style = 'Table Grid'
    
    # Заполняем таблицу
    table.cell(0, 0).text = 'Продукт'
    table.cell(0, 1).text = 'Количество'
    table.cell(0, 2).text = 'Сумма'
    
    table.cell(1, 0).text = 'Товар А'
    table.cell(1, 1).text = '12'  # Изменено
    table.cell(1, 2).text = '60 000'  # Изменено
    
    table.cell(2, 0).text = 'Товар Б'
    table.cell(2, 1).text = '5'
    table.cell(2, 2).text = '50 000'
    
    table.cell(3, 0).text = 'Товар В'  # Добавлено
    table.cell(3, 1).text = '2'  # Добавлено
    table.cell(3, 2).text = '10 000'  # Добавлено
    
    # Добавляем еще один раздел (изменен)
    doc.add_heading('Раздел 2: Выводы', level=1)
    doc.add_paragraph('Продажи превысили план на 20%.')  # Изменено
    
    # Добавляем новый раздел
    doc.add_heading('Раздел 3: Рекомендации', level=1)
    doc.add_paragraph('Рекомендуется увеличить производство товара А.')  # Добавлено
    
    # Сохраняем в байты
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def test_auto_analysis():
    """Тестирование автоматического анализа"""
    print("🧪 Тестирование автоматического анализа документов")
    print("=" * 60)
    
    # Получаем или создаем тестового пользователя
    user, created = User.objects.get_or_create(
        username='testuser',
        defaults={
            'email': 'test@example.com',
            'first_name': 'Тест',
            'last_name': 'Пользователь',
            'role': 'user'
        }
    )
    if created:
        user.set_password('testpass123')
        user.save()
        print(f"    ✅ Создан пользователь: {user.username}")
    else:
        print(f"    ✅ Используется существующий пользователь: {user.username}")
    
    # Создаем тестовые .docx файлы
    print("\n📄 Создание тестовых документов...")
    docx_v1_content = create_test_docx_v1()
    docx_v2_content = create_test_docx_v2()
    
    test_file_v1 = SimpleUploadedFile(
        "test_report_v1.docx",
        docx_v1_content,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    test_file_v2 = SimpleUploadedFile(
        "test_report_v2.docx",
        docx_v2_content,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    
    print(f"    ✅ Создан файл v1: {test_file_v1.name} ({len(docx_v1_content)} байт)")
    print(f"    ✅ Создан файл v2: {test_file_v2.name} ({len(docx_v2_content)} байт)")
    
    # Создаем клиент для тестирования
    client = Client()
    client.force_login(user)
    
    # Тест 1: Загрузка документов
    print("\n🔍 Тест 1: Загрузка документов")
    
    # Загружаем первый документ
    response1 = client.post('/upload/', {
        'title': 'Отчет о продажах v1',
        'file': test_file_v1
    })
    
    if response1.status_code == 302:
        print("    ✅ Первый документ загружен")
    else:
        print(f"    ❌ Ошибка загрузки первого документа: {response1.status_code}")
        return
    
    # Загружаем второй документ
    response2 = client.post('/upload/', {
        'title': 'Отчет о продажах v2',
        'file': test_file_v2
    })
    
    if response2.status_code == 302:
        print("    ✅ Второй документ загружен")
    else:
        print(f"    ❌ Ошибка загрузки второго документа: {response2.status_code}")
        return
    
    # Получаем загруженные документы
    documents = DjangoDocument.objects.filter(user=user).order_by('-upload_date')[:2]
    if len(documents) < 2:
        print("    ❌ Недостаточно документов для сравнения")
        return
    
    doc1, doc2 = documents[0], documents[1]
    print(f"    📝 Документ 1: {doc1.title} (ID: {doc1.id})")
    print(f"    📝 Документ 2: {doc2.title} (ID: {doc2.id})")
    
    # Тест 2: Создание сравнения через форму
    print("\n🔍 Тест 2: Создание сравнения через форму")
    
    response = client.post('/analysis/create/', {
        'title': 'Автоматический анализ тест',
        'base_document': doc1.id,
        'compared_document': doc2.id
    })
    
    if response.status_code == 302:  # Редирект после успешного создания
        print("    ✅ Сравнение успешно создано")
        
        # Получаем последнее созданное сравнение
        latest_comparison = Comparison.objects.filter(user=user).order_by('-created_date').first()
        if latest_comparison:
            print(f"    📝 Сравнение ID: {latest_comparison.id}")
            print(f"    📝 Название: {latest_comparison.title}")
            print(f"    📝 Статус: {latest_comparison.status}")
            
            # Проверяем, был ли анализ выполнен автоматически
            if latest_comparison.status == 'completed':
                print("    ✅ Анализ автоматически завершен!")
                
                # Проверяем результаты анализа
                if latest_comparison.changes_summary:
                    print(f"    📊 Сводка изменений: {latest_comparison.changes_summary}")
                
                # Проверяем найденные изменения
                changes = latest_comparison.changes.all()
                print(f"    🔍 Найдено изменений: {changes.count()}")
                
                for i, change in enumerate(changes[:5], 1):  # Показываем первые 5
                    print(f"        {i}. {change.change_type}: {change.section} - {change.old_value[:30]}... -> {change.new_value[:30]}...")
                
                if changes.count() > 0:
                    print("    ✅ Автоматический анализ работает корректно!")
                else:
                    print("    ⚠️ Анализ выполнен, но изменения не найдены")
            else:
                print(f"    ❌ Анализ не был автоматически завершен (статус: {latest_comparison.status})")
        else:
            print("    ❌ Сравнение не найдено в базе данных")
    else:
        print(f"    ❌ Ошибка создания сравнения: статус {response.status_code}")
        if hasattr(response, 'content'):
            print(f"    📄 Содержимое ответа: {response.content.decode()[:200]}...")
    
    # Тест 3: Проверка ручного анализа (если автоматический не сработал)
    print("\n🔍 Тест 3: Проверка ручного анализа")
    latest_comparison = Comparison.objects.filter(user=user).order_by('-created_date').first()
    if latest_comparison and latest_comparison.status != 'completed':
        print(f"    🔧 Запуск ручного анализа сравнения {latest_comparison.id}...")
        
        try:
            comparison_service = DocumentComparisonService()
            analysis_result = comparison_service.compare_documents(latest_comparison)
            
            if analysis_result and 'changes' in analysis_result:
                changes_count = len(analysis_result['changes'])
                print("    ✅ Ручной анализ успешен!")
                print(f"    🔍 Найдено изменений: {changes_count}")
                
                # Сохраняем результаты
                comparison_service.save_comparison_results(latest_comparison, analysis_result)
                latest_comparison.status = 'completed'
                latest_comparison.save()
                print("    ✅ Результаты сохранены")
            else:
                print("    ❌ Ошибка ручного анализа: нет данных в результате")
        except Exception as e:
            print(f"    ❌ Исключение при ручном анализе: {e}")
    else:
        print("    ⏭️ Ручной анализ не требуется")
    
    # Тест 4: Проверка интерфейса
    print("\n🔍 Тест 4: Проверка интерфейса")
    response = client.get('/analysis/')
    if response.status_code == 200:
        print("    ✅ Страница анализа загружена")
        
        content = response.content.decode()
        
        # Проверяем наличие информации об автоматическом анализе
        if 'Автоматический анализ' in content:
            print("    ✅ Информация об автоматическом анализе найдена")
        else:
            print("    ⚠️ Информация об автоматическом анализе не найдена")
        
        # Проверяем кнопки действий
        if 'fas fa-play' in content:
            print("    ✅ Кнопки анализа найдены")
        else:
            print("    ⚠️ Кнопки анализа не найдены")
    else:
        print(f"    ❌ Ошибка загрузки страницы: {response.status_code}")
    
    # Тест 5: Проверка страницы создания сравнения
    print("\n🔍 Тест 5: Проверка страницы создания сравнения")
    response = client.get('/analysis/create/')
    if response.status_code == 200:
        print("    ✅ Страница создания сравнения доступна")
        
        content = response.content.decode()
        if 'Автоматический анализ' in content:
            print("    ✅ Информация об автоматическом анализе на странице создания")
        else:
            print("    ❌ Информация об автоматическом анализе отсутствует на странице создания")
    else:
        print(f"    ❌ Страница создания сравнения недоступна: {response.status_code}")
    
    print("\n🎉 Тестирование автоматического анализа завершено!")
    print("=" * 60)
    
    # Итоговая статистика
    total_comparisons = Comparison.objects.filter(user=user).count()
    completed_comparisons = Comparison.objects.filter(user=user, status='completed').count()
    
    print(f"📊 Итоговая статистика:")
    print(f"    🔍 Всего сравнений: {total_comparisons}")
    print(f"    ✅ Завершенных анализов: {completed_comparisons}")
    print(f"    📈 Процент завершения: {(completed_comparisons/total_comparisons*100):.1f}%" if total_comparisons > 0 else "    📈 Процент завершения: 0%")

def main():
    """Основная функция тестирования"""
    try:
        test_auto_analysis()
        return True
    except Exception as e:
        print(f"\n❌ Ошибка во время тестирования: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)
