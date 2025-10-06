"""
–°–µ—Ä–≤–∏—Å—ã –¥–ª—è —Ä–∞–±–æ—Ç—ã —Å –¥–æ–∫—É–º–µ–Ω—Ç–∞–º–∏ Word
"""
import os
import logging
from typing import Dict, List, Any, Optional
from docx import Document as DocxDocument
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph
from django.core.files.uploadedfile import UploadedFile
from django.core.exceptions import ValidationError
from .models import Document, DocumentSection, DocumentTable, DocumentTableAnalysis

logger = logging.getLogger(__name__)


class DocumentParserService:
    """
    –°–µ—Ä–≤–∏—Å –¥–ª—è –ø–∞—Ä—Å–∏–Ω–≥–∞ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤ Word (.docx)
    """
    
    def __init__(self):
        self.supported_formats = ['.docx']
        self.max_file_size = 10 * 1024 * 1024  # 10 –ú–ë
    
    def validate_file(self, file: UploadedFile) -> bool:
        """
        –í–∞–ª–∏–¥–∞—Ü–∏—è –∑–∞–≥—Ä—É–∂–µ–Ω–Ω–æ–≥–æ —Ñ–∞–π–ª–∞
        """
        # –ü—Ä–æ–≤–µ—Ä–∫–∞ —Ä–∞–∑–º–µ—Ä–∞
        if file.size > self.max_file_size:
            raise ValidationError(f"–†–∞–∑–º–µ—Ä —Ñ–∞–π–ª–∞ –ø—Ä–µ–≤—ã—à–∞–µ—Ç {self.max_file_size // (1024*1024)} –ú–ë")
        
        # –ü—Ä–æ–≤–µ—Ä–∫–∞ —Ä–∞—Å—à–∏—Ä–µ–Ω–∏—è
        file_extension = os.path.splitext(file.name)[1].lower()
        if file_extension not in self.supported_formats:
            raise ValidationError(f"–ù–µ–ø–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ–º—ã–π —Ñ–æ—Ä–º–∞—Ç —Ñ–∞–π–ª–∞. –ü–æ–¥–¥–µ—Ä–∂–∏–≤–∞—é—Ç—Å—è: {', '.join(self.supported_formats)}")
        
        return True
    
    def parse_document(self, document: Document) -> Dict[str, Any]:
        """
        –ü–∞—Ä—Å–∏–Ω–≥ –¥–æ–∫—É–º–µ–Ω—Ç–∞ –∏ –∏–∑–≤–ª–µ—á–µ–Ω–∏–µ —Å–æ–¥–µ—Ä–∂–∏–º–æ–≥–æ
        """
        try:
            # –û—Ç–∫—Ä—ã–≤–∞–µ–º –¥–æ–∫—É–º–µ–Ω—Ç
            docx_doc = DocxDocument(document.file.path)
            
            # –ò–∑–≤–ª–µ–∫–∞–µ–º —Å–æ–¥–µ—Ä–∂–∏–º–æ–µ
            content_data = {
                'text_content': self._extract_text(docx_doc),
                'sections': self._extract_sections(docx_doc),
                'tables': self._extract_tables(docx_doc),
                'metadata': self._extract_metadata(docx_doc, document),
                'structure': self._analyze_structure(docx_doc)
            }
            
            # –°–æ—Ö—Ä–∞–Ω—è–µ–º –∏–∑–≤–ª–µ—á–µ–Ω–Ω—ã–π —Ç–µ–∫—Å—Ç –≤ –º–æ–¥–µ–ª—å
            document.content_text = content_data['text_content']
            document.save(update_fields=['content_text'])
            
            # –ê–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏ –∞–Ω–∞–ª–∏–∑–∏—Ä—É–µ–º —Ç–∞–±–ª–∏—Ü—ã, –µ—Å–ª–∏ –æ–Ω–∏ –µ—Å—Ç—å
            tables_count = len(content_data.get('tables', []))
            if tables_count > 0:
                logger.info(f"–ù–∞—á–∏–Ω–∞–µ–º –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏–π –∞–Ω–∞–ª–∏–∑ —Ç–∞–±–ª–∏—Ü –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.title}: –Ω–∞–π–¥–µ–Ω–æ {tables_count} —Ç–∞–±–ª–∏—Ü")
                try:
                    table_analysis_service = DocumentTableAnalysisService()
                    analysis_result = table_analysis_service.analyze_document_tables(document)
                    if analysis_result['success']:
                        analyzed_count = analysis_result.get('tables_count', 0)
                        logger.info(f"‚úÖ –ê–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏–π –∞–Ω–∞–ª–∏–∑ —Ç–∞–±–ª–∏—Ü —É—Å–ø–µ—à–Ω–æ –≤—ã–ø–æ–ª–Ω–µ–Ω –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.title}: –ø—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–æ {analyzed_count} –∏–∑ {tables_count} —Ç–∞–±–ª–∏—Ü")
                        
                        # –õ–æ–≥–∏—Ä—É–µ–º —Å–≤–æ–¥–∫—É –∞–Ω–∞–ª–∏–∑–∞
                        summary = analysis_result.get('summary', {})
                        if summary:
                            logger.info(f"üìä –°–≤–æ–¥–∫–∞ –∞–Ω–∞–ª–∏–∑–∞ —Ç–∞–±–ª–∏—Ü: {summary.get('total_cells', 0)} —è—á–µ–µ–∫, –∑–∞–ø–æ–ª–Ω–µ–Ω–Ω–æ—Å—Ç—å {summary.get('fill_percentage', 0)}%")
                    else:
                        logger.warning(f"‚ö†Ô∏è –û—à–∏–±–∫–∞ –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–æ–≥–æ –∞–Ω–∞–ª–∏–∑–∞ —Ç–∞–±–ª–∏—Ü –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.title}: {analysis_result.get('error', '–ù–µ–∏–∑–≤–µ—Å—Ç–Ω–∞—è –æ—à–∏–±–∫–∞')}")
                except Exception as e:
                    logger.error(f"‚ùå –ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ –ø—Ä–∏ –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–æ–º –∞–Ω–∞–ª–∏–∑–µ —Ç–∞–±–ª–∏—Ü –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.title}: {str(e)}")
            else:
                logger.info(f"üìÑ –î–æ–∫—É–º–µ–Ω—Ç {document.title} –Ω–µ —Å–æ–¥–µ—Ä–∂–∏—Ç —Ç–∞–±–ª–∏—Ü - –∞–Ω–∞–ª–∏–∑ –Ω–µ —Ç—Ä–µ–±—É–µ—Ç—Å—è")
            
            # –ê–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏ –≥–µ–Ω–µ—Ä–∏—Ä—É–µ–º –∫–ª—é—á–µ–≤—ã–µ –º–æ–º–µ–Ω—Ç—ã
            logger.info(f"–ù–∞—á–∏–Ω–∞–µ–º –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫—É—é –≥–µ–Ω–µ—Ä–∞—Ü–∏—é –∫–ª—é—á–µ–≤—ã—Ö –º–æ–º–µ–Ω—Ç–æ–≤ –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.title}")
            try:
                # –ò–º–ø–æ—Ä—Ç–∏—Ä—É–µ–º DocumentKeyPointsService –ª–æ–∫–∞–ª—å–Ω–æ –¥–ª—è –∏–∑–±–µ–∂–∞–Ω–∏—è —Ü–∏–∫–ª–∏—á–µ—Å–∫–∏—Ö –∏–º–ø–æ—Ä—Ç–æ–≤
                from documents.services import DocumentKeyPointsService
                key_points_service = DocumentKeyPointsService()
                key_points_result = key_points_service.generate_key_points(document)
                
                if key_points_result['success']:
                    key_points_count = len(key_points_result.get('key_points', []))
                    logger.info(f"‚úÖ –ê–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∞—è –≥–µ–Ω–µ—Ä–∞—Ü–∏—è –∫–ª—é—á–µ–≤—ã—Ö –º–æ–º–µ–Ω—Ç–æ–≤ —É—Å–ø–µ—à–Ω–æ –≤—ã–ø–æ–ª–Ω–µ–Ω–∞ –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.title}: —Å–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–Ω–æ {key_points_count} –º–æ–º–µ–Ω—Ç–æ–≤")
                    
                    # –õ–æ–≥–∏—Ä—É–µ–º –∫—Ä–∞—Ç–∫—É—é –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –æ –∫–ª—é—á–µ–≤—ã—Ö –º–æ–º–µ–Ω—Ç–∞—Ö
                    if key_points_result.get('summary'):
                        logger.info(f"üìù –ö—Ä–∞—Ç–∫–æ–µ —Ä–µ–∑—é–º–µ: {key_points_result['summary'][:100]}...")
                else:
                    logger.warning(f"‚ö†Ô∏è –û—à–∏–±–∫–∞ –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–æ–π –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –∫–ª—é—á–µ–≤—ã—Ö –º–æ–º–µ–Ω—Ç–æ–≤ –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.title}: {key_points_result.get('error', '–ù–µ–∏–∑–≤–µ—Å—Ç–Ω–∞—è –æ—à–∏–±–∫–∞')}")
            except Exception as e:
                logger.error(f"‚ùå –ö—Ä–∏—Ç–∏—á–µ—Å–∫–∞—è –æ—à–∏–±–∫–∞ –ø—Ä–∏ –∞–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–æ–π –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –∫–ª—é—á–µ–≤—ã—Ö –º–æ–º–µ–Ω—Ç–æ–≤ –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.title}: {str(e)}")
            
            logger.info(f"–î–æ–∫—É–º–µ–Ω—Ç {document.title} —É—Å–ø–µ—à–Ω–æ –æ–±—Ä–∞–±–æ—Ç–∞–Ω")
            
            # –î–æ–±–∞–≤–ª—è–µ–º –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –æ–± –∞–Ω–∞–ª–∏–∑–µ —Ç–∞–±–ª–∏—Ü –∏ –∫–ª—é—á–µ–≤—ã—Ö –º–æ–º–µ–Ω—Ç–æ–≤ –≤ —Ä–µ–∑—É–ª—å—Ç–∞—Ç
            result_data = {
                "success": True,
                "content_data": content_data
            }
            
            if tables_count > 0:
                result_data["table_analysis"] = {
                    "performed": True,
                    "tables_found": tables_count,
                    "message": f"–ê–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏ –ø—Ä–æ–∞–Ω–∞–ª–∏–∑–∏—Ä–æ–≤–∞–Ω–æ {tables_count} —Ç–∞–±–ª–∏—Ü"
                }
            else:
                result_data["table_analysis"] = {
                    "performed": False,
                    "tables_found": 0,
                    "message": "–¢–∞–±–ª–∏—Ü—ã –Ω–µ –Ω–∞–π–¥–µ–Ω—ã"
                }
            
            # –î–æ–±–∞–≤–ª—è–µ–º –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏—é –æ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –∫–ª—é—á–µ–≤—ã—Ö –º–æ–º–µ–Ω—Ç–æ–≤
            result_data["key_points_generation"] = {
                "performed": True,
                "message": "–ê–≤—Ç–æ–º–∞—Ç–∏—á–µ—Å–∫–∏ —Å–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–Ω—ã –∫–ª—é—á–µ–≤—ã–µ –º–æ–º–µ–Ω—Ç—ã"
            }
            
            return result_data
            
        except Exception as e:
            logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –ø–∞—Ä—Å–∏–Ω–≥–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.title}: {str(e)}")
            return {
                "success": False,
                "error": f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –æ–±—Ä–∞–±–æ—Ç–∫–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞: {str(e)}"
            }
    
    def save_parsed_content(self, document: Document, content_data: Dict[str, Any]) -> None:
        """
        –°–æ—Ö—Ä–∞–Ω—è–µ—Ç —Ä–∞—Å–ø–∞—Ä—Å–µ–Ω–Ω–æ–µ —Å–æ–¥–µ—Ä–∂–∏–º–æ–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞ –≤ –±–∞–∑—É –¥–∞–Ω–Ω—ã—Ö
        """
        try:
            # –£–¥–∞–ª—è–µ–º —Å—Ç–∞—Ä—ã–µ –¥–∞–Ω–Ω—ã–µ
            DocumentSection.objects.filter(document=document).delete()
            DocumentTable.objects.filter(document=document).delete()
            
            # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ä–∞–∑–¥–µ–ª—ã
            for section_data in content_data.get('sections', []):
                DocumentSection.objects.create(
                    document=document,
                    title=section_data['title'],
                    content=section_data['content'],
                    level=section_data['level'],
                    order=section_data['order']
                )
            
            # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–∞–±–ª–∏—Ü—ã
            for table_data in content_data.get('tables', []):
                DocumentTable.objects.create(
                    document=document,
                    title=table_data['title'],
                    content=table_data['content'],
                    rows=table_data['rows'],
                    columns=table_data['columns'],
                    order=table_data['order']
                )
            
            logger.info(f"–°–æ–¥–µ—Ä–∂–∏–º–æ–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.title} —Å–æ—Ö—Ä–∞–Ω–µ–Ω–æ –≤ –ë–î")
            
        except Exception as e:
            logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–∏ —Å–æ–¥–µ—Ä–∂–∏–º–æ–≥–æ –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.title}: {str(e)}")
            raise
    
    def _extract_text(self, doc: DocumentType) -> str:
        """
        –ò–∑–≤–ª–µ—á–µ–Ω–∏–µ –≤—Å–µ–≥–æ —Ç–µ–∫—Å—Ç–æ–≤–æ–≥–æ —Å–æ–¥–µ—Ä–∂–∏–º–æ–≥–æ –≤–∫–ª—é—á–∞—è —Ç–∞–±–ª–∏—Ü—ã
        """
        full_text = []
        
        # –ò–∑–≤–ª–µ–∫–∞–µ–º —Ç–µ–∫—Å—Ç –∏–∑ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–æ–≤
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        
        # –ò–∑–≤–ª–µ–∫–∞–µ–º —Ç–µ–∫—Å—Ç –∏–∑ —Ç–∞–±–ª–∏—Ü
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    table_text.append(' | '.join(row_text))
            
            if table_text:
                full_text.append('\n'.join(table_text))
        
        return '\n\n'.join(full_text)
    
    def _extract_sections(self, doc: DocumentType) -> List[Dict[str, Any]]:
        """
        –ò–∑–≤–ª–µ—á–µ–Ω–∏–µ —Ä–∞–∑–¥–µ–ª–æ–≤ –¥–æ–∫—É–º–µ–Ω—Ç–∞
        """
        sections = []
        current_section = None
        order = 0
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            
            # –û–ø—Ä–µ–¥–µ–ª—è–µ–º —É—Ä–æ–≤–µ–Ω—å –∑–∞–≥–æ–ª–æ–≤–∫–∞ –ø–æ —Å—Ç–∏–ª—é
            level = self._get_heading_level(paragraph)
            
            if level > 0:  # –≠—Ç–æ –∑–∞–≥–æ–ª–æ–≤–æ–∫
                # –°–æ—Ö—Ä–∞–Ω—è–µ–º –ø—Ä–µ–¥—ã–¥—É—â–∏–π —Ä–∞–∑–¥–µ–ª
                if current_section:
                    sections.append(current_section)
                
                # –°–æ–∑–¥–∞–µ–º –Ω–æ–≤—ã–π —Ä–∞–∑–¥–µ–ª
                current_section = {
                    'title': text,
                    'level': level,
                    'order': order,
                    'content': '',
                    'paragraphs': []
                }
                order += 1
            else:
                # –î–æ–±–∞–≤–ª—è–µ–º –∞–±–∑–∞—Ü –∫ —Ç–µ–∫—É—â–µ–º—É —Ä–∞–∑–¥–µ–ª—É
                if current_section:
                    current_section['content'] += text + '\n\n'
                    current_section['paragraphs'].append(text)
                else:
                    # –ï—Å–ª–∏ –Ω–µ—Ç —Ç–µ–∫—É—â–µ–≥–æ —Ä–∞–∑–¥–µ–ª–∞, —Å–æ–∑–¥–∞–µ–º –æ–±—â–∏–π
                    if not current_section:
                        current_section = {
                            'title': '–û–±—â–∏–π —Ç–µ–∫—Å—Ç',
                            'level': 0,
                            'order': order,
                            'content': text + '\n\n',
                            'paragraphs': [text]
                        }
                    else:
                        current_section['content'] += text + '\n\n'
                        current_section['paragraphs'].append(text)
        
        # –î–æ–±–∞–≤–ª—è–µ–º –ø–æ—Å–ª–µ–¥–Ω–∏–π —Ä–∞–∑–¥–µ–ª
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def _extract_tables(self, doc: DocumentType) -> List[Dict[str, Any]]:
        """
        –ò–∑–≤–ª–µ—á–µ–Ω–∏–µ —Ç–∞–±–ª–∏—Ü –∏–∑ –¥–æ–∫—É–º–µ–Ω—Ç–∞
        """
        tables = []
        
        for i, table in enumerate(doc.tables):
            table_data = {
                'order': i,
                'title': f'–¢–∞–±–ª–∏—Ü–∞ {i + 1}',
                'rows': [],
                'headers': [],
                'row_count': len(table.rows),
                'col_count': len(table.columns) if table.rows else 0
            }
            
            # –ò–∑–≤–ª–µ–∫–∞–µ–º –¥–∞–Ω–Ω—ã–µ —Ç–∞–±–ª–∏—Ü—ã
            for row_idx, row in enumerate(table.rows):
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    row_data.append(cell_text)
                
                table_data['rows'].append(row_data)
                
                # –ü–µ—Ä–≤–∞—è —Å—Ç—Ä–æ–∫–∞ —Å—á–∏—Ç–∞–µ—Ç—Å—è –∑–∞–≥–æ–ª–æ–≤–∫–∞–º–∏
                if row_idx == 0:
                    table_data['headers'] = row_data
            
            tables.append(table_data)
        
        return tables
    
    def _extract_metadata(self, doc: DocumentType, document: Document) -> Dict[str, Any]:
        """
        –ò–∑–≤–ª–µ—á–µ–Ω–∏–µ –º–µ—Ç–∞–¥–∞–Ω–Ω—ã—Ö –¥–æ–∫—É–º–µ–Ω—Ç–∞
        """
        core_props = doc.core_properties
        
        metadata = {
            'title': core_props.title or document.title,
            'author': core_props.author or '',
            'created': core_props.created.isoformat() if core_props.created else None,
            'modified': core_props.modified.isoformat() if core_props.modified else None,
            'subject': core_props.subject or '',
            'keywords': core_props.keywords or '',
            'comments': core_props.comments or '',
            'word_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
            'page_count': self._estimate_page_count(doc)
        }
        
        return metadata
    
    def _analyze_structure(self, doc: DocumentType) -> Dict[str, Any]:
        """
        –ê–Ω–∞–ª–∏–∑ —Å—Ç—Ä—É–∫—Ç—É—Ä—ã –¥–æ–∫—É–º–µ–Ω—Ç–∞
        """
        structure = {
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables),
            'heading_levels': set(),
            'has_tables': len(doc.tables) > 0,
            'has_images': False,  # –ü–æ–∫–∞ –Ω–µ —Ä–µ–∞–ª–∏–∑–æ–≤–∞–Ω–æ
            'estimated_pages': self._estimate_page_count(doc)
        }
        
        # –ê–Ω–∞–ª–∏–∑–∏—Ä—É–µ–º —É—Ä–æ–≤–Ω–∏ –∑–∞–≥–æ–ª–æ–≤–∫–æ–≤
        for paragraph in doc.paragraphs:
            level = self._get_heading_level(paragraph)
            if level > 0:
                structure['heading_levels'].add(level)
        
        structure['heading_levels'] = sorted(list(structure['heading_levels']))
        
        return structure
    
    def _get_heading_level(self, paragraph: Paragraph) -> int:
        """
        –û–ø—Ä–µ–¥–µ–ª–µ–Ω–∏–µ —É—Ä–æ–≤–Ω—è –∑–∞–≥–æ–ª–æ–≤–∫–∞
        """
        style_name = paragraph.style.name.lower()
        
        if 'heading 1' in style_name or '–∑–∞–≥–æ–ª–æ–≤–æ–∫ 1' in style_name:
            return 1
        elif 'heading 2' in style_name or '–∑–∞–≥–æ–ª–æ–≤–æ–∫ 2' in style_name:
            return 2
        elif 'heading 3' in style_name or '–∑–∞–≥–æ–ª–æ–≤–æ–∫ 3' in style_name:
            return 3
        elif 'heading 4' in style_name or '–∑–∞–≥–æ–ª–æ–≤–æ–∫ 4' in style_name:
            return 4
        elif 'heading 5' in style_name or '–∑–∞–≥–æ–ª–æ–≤–æ–∫ 5' in style_name:
            return 5
        elif 'heading 6' in style_name or '–∑–∞–≥–æ–ª–æ–≤–æ–∫ 6' in style_name:
            return 6
        else:
            return 0
    
    def _estimate_page_count(self, doc: DocumentType) -> int:
        """
        –ü—Ä–∏–º–µ—Ä–Ω–∞—è –æ—Ü–µ–Ω–∫–∞ –∫–æ–ª–∏—á–µ—Å—Ç–≤–∞ —Å—Ç—Ä–∞–Ω–∏—Ü
        """
        # –ü—Ä–æ—Å—Ç–∞—è –æ—Ü–µ–Ω–∫–∞: –ø—Ä–∏–º–µ—Ä–Ω–æ 500 —Å–∏–º–≤–æ–ª–æ–≤ –Ω–∞ —Å—Ç—Ä–∞–Ω–∏—Ü—É
        total_chars = sum(len(p.text) for p in doc.paragraphs)
        estimated_pages = max(1, total_chars // 500)
        return estimated_pages
    
    def save_parsed_content(self, document: Document, content_data: Dict[str, Any]) -> None:
        """
        –°–æ—Ö—Ä–∞–Ω–µ–Ω–∏–µ –∏–∑–≤–ª–µ—á–µ–Ω–Ω–æ–≥–æ —Å–æ–¥–µ—Ä–∂–∏–º–æ–≥–æ –≤ –±–∞–∑—É –¥–∞–Ω–Ω—ã—Ö
        """
        try:
            # –û–±–Ω–æ–≤–ª—è–µ–º –æ—Å–Ω–æ–≤–Ω–æ–π –¥–æ–∫—É–º–µ–Ω—Ç
            document.content_text = content_data['text_content']
            document.content_structure = content_data['structure']
            document.metadata = content_data['metadata']
            document.status = 'processed'
            document.save()
            
            # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ä–∞–∑–¥–µ–ª—ã
            DocumentSection.objects.filter(document=document).delete()
            for section_data in content_data['sections']:
                DocumentSection.objects.create(
                    document=document,
                    title=section_data['title'],
                    content=section_data['content'],
                    order=section_data['order'],
                    level=section_data['level']
                )
            
            # –°–æ—Ö—Ä–∞–Ω—è–µ–º —Ç–∞–±–ª–∏—Ü—ã
            DocumentTable.objects.filter(document=document).delete()
            for table_data in content_data['tables']:
                DocumentTable.objects.create(
                    document=document,
                    title=table_data['title'],
                    data=table_data,
                    order=table_data['order']
                )
            
            logger.info(f"–°–æ–¥–µ—Ä–∂–∏–º–æ–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.title} —Å–æ—Ö—Ä–∞–Ω–µ–Ω–æ –≤ –ë–î")
            
        except Exception as e:
            logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–∏ —Å–æ–¥–µ—Ä–∂–∏–º–æ–≥–æ –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.title}: {str(e)}")
            raise


class DocumentValidationService:
    """
    –°–µ—Ä–≤–∏—Å –¥–ª—è –≤–∞–ª–∏–¥–∞—Ü–∏–∏ –¥–æ–∫—É–º–µ–Ω—Ç–æ–≤
    """
    
    def __init__(self):
        self.allowed_extensions = ['.docx']
        self.max_file_size = 10 * 1024 * 1024  # 10 –ú–ë
    
    def validate_upload(self, file: UploadedFile) -> Dict[str, Any]:
        """
        –ü–æ–ª–Ω–∞—è –≤–∞–ª–∏–¥–∞—Ü–∏—è –∑–∞–≥—Ä—É–∂–µ–Ω–Ω–æ–≥–æ —Ñ–∞–π–ª–∞
        """
        validation_result = {
            'is_valid': True,
            'errors': [],
            'warnings': [],
            'file_info': {}
        }
        
        # –ü—Ä–æ–≤–µ—Ä–∫–∞ —Ä–∞–∑–º–µ—Ä–∞
        if file.size > self.max_file_size:
            validation_result['is_valid'] = False
            validation_result['errors'].append(
                f"–†–∞–∑–º–µ—Ä —Ñ–∞–π–ª–∞ ({file.size // (1024*1024)} –ú–ë) –ø—Ä–µ–≤—ã—à–∞–µ—Ç –º–∞–∫—Å–∏–º–∞–ª—å–Ω–æ –¥–æ–ø—É—Å—Ç–∏–º—ã–π ({self.max_file_size // (1024*1024)} –ú–ë)"
            )
        
        # –ü—Ä–æ–≤–µ—Ä–∫–∞ —Ä–∞—Å—à–∏—Ä–µ–Ω–∏—è
        file_extension = os.path.splitext(file.name)[1].lower()
        if file_extension not in self.allowed_extensions:
            validation_result['is_valid'] = False
            validation_result['errors'].append(
                f"–ù–µ–ø–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ–º—ã–π —Ñ–æ—Ä–º–∞—Ç —Ñ–∞–π–ª–∞. –ü–æ–¥–¥–µ—Ä–∂–∏–≤–∞—é—Ç—Å—è: {', '.join(self.allowed_extensions)}"
            )
        
        # –ü—Ä–æ–≤–µ—Ä–∫–∞ –∏–º–µ–Ω–∏ —Ñ–∞–π–ª–∞
        if len(file.name) > 255:
            validation_result['is_valid'] = False
            validation_result['errors'].append("–ò–º—è —Ñ–∞–π–ª–∞ —Å–ª–∏—à–∫–æ–º –¥–ª–∏–Ω–Ω–æ–µ (–º–∞–∫—Å–∏–º—É–º 255 —Å–∏–º–≤–æ–ª–æ–≤)")
        
        # –ò–Ω—Ñ–æ—Ä–º–∞—Ü–∏—è –æ —Ñ–∞–π–ª–µ
        validation_result['file_info'] = {
            'name': file.name,
            'size': file.size,
            'size_mb': round(file.size / (1024 * 1024), 2),
            'extension': file_extension,
            'content_type': getattr(file, 'content_type', 'application/octet-stream')
        }
        
        return validation_result


class DocumentKeyPointsService:
    """
    –°–µ—Ä–≤–∏—Å –¥–ª—è –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –∫–ª—é—á–µ–≤—ã—Ö –º–æ–º–µ–Ω—Ç–æ–≤ –¥–æ–∫—É–º–µ–Ω—Ç–∞
    """
    
    def __init__(self):
        from analysis.ollama_service import OllamaService
        from django.utils import timezone
        from settings.models import ApplicationSettings
        
        # –ü–æ–ª—É—á–∞–µ–º –º–æ–¥–µ–ª—å –∏–∑ –Ω–∞—Å—Ç—Ä–æ–µ–∫ –ø—Ä–∏–ª–æ–∂–µ–Ω–∏—è
        settings = ApplicationSettings.get_settings()
        model = settings.default_neural_network_model or 'llama3'
        
        self.ollama_service = OllamaService(model=model)
        self.timezone = timezone
    
    def _get_table_rows_count(self, document):
        """
        –ü–æ–ª—É—á–∞–µ—Ç –æ–±—â–µ–µ –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ —Å—Ç—Ä–æ–∫ –≤–æ –≤—Å–µ—Ö —Ç–∞–±–ª–∏—Ü–∞—Ö –¥–æ–∫—É–º–µ–Ω—Ç–∞
        
        Args:
            document: –û–±—ä–µ–∫—Ç Document
            
        Returns:
            int: –û–±—â–µ–µ –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ —Å—Ç—Ä–æ–∫ –≤ —Ç–∞–±–ª–∏—Ü–∞—Ö
        """
        try:
            # –ü–æ–ª—É—á–∞–µ–º –≤—Å–µ —Ç–∞–±–ª–∏—Ü—ã –¥–æ–∫—É–º–µ–Ω—Ç–∞
            tables = document.tables.all()
            
            if not tables.exists():
                return 0
            
            # –°—á–∏—Ç–∞–µ–º –æ–±—â–µ–µ –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ —Å—Ç—Ä–æ–∫
            total_rows = 0
            for table in tables:
                # –ü–æ–ª—É—á–∞–µ–º –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ —Å—Ç—Ä–æ–∫ –∏–∑ –¥–∞–Ω–Ω—ã—Ö —Ç–∞–±–ª–∏—Ü—ã
                table_data = table.data
                if table_data and isinstance(table_data, dict):
                    # –ò—Å–ø–æ–ª—å–∑—É–µ–º row_count –∏–∑ –¥–∞–Ω–Ω—ã—Ö —Ç–∞–±–ª–∏—Ü—ã
                    total_rows += table_data.get('row_count', 0)
                elif table_data and isinstance(table_data, list):
                    total_rows += len(table_data)
                else:
                    # –ï—Å–ª–∏ —Å—Ç—Ä—É–∫—Ç—É—Ä–∞ –Ω–µ–∏–∑–≤–µ—Å—Ç–Ω–∞, –æ—Ü–µ–Ω–∏–≤–∞–µ–º –ø–æ —Ä–∞–∑–º–µ—Ä—É
                    total_rows += 5  # –ü—Ä–∏–º–µ—Ä–Ω–∞—è –æ—Ü–µ–Ω–∫–∞
            
            logger.info(f"–î–æ–∫—É–º–µ–Ω—Ç {document.id}: –Ω–∞–π–¥–µ–Ω–æ {tables.count()} —Ç–∞–±–ª–∏—Ü —Å –æ–±—â–∏–º –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ–º —Å—Ç—Ä–æ–∫ {total_rows}")
            return total_rows
            
        except Exception as e:
            logger.warning(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –ø–æ–¥—Å—á–µ—Ç–µ —Å—Ç—Ä–æ–∫ —Ç–∞–±–ª–∏—Ü –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.id}: {str(e)}")
            return 0
    
    def generate_key_points(self, document):
        """
        –ì–µ–Ω–µ—Ä–∏—Ä—É–µ—Ç –∫–ª—é—á–µ–≤—ã–µ –º–æ–º–µ–Ω—Ç—ã –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞
        
        Args:
            document: –û–±—ä–µ–∫—Ç Document
            
        Returns:
            Dict —Å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞–º–∏ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏
        """
        try:
            if not document.has_content():
                raise ValueError("–î–æ–∫—É–º–µ–Ω—Ç –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å –æ–±—Ä–∞–±–æ—Ç–∞–Ω –ø–µ—Ä–µ–¥ –≥–µ–Ω–µ—Ä–∞—Ü–∏–µ–π –∫–ª—é—á–µ–≤—ã—Ö –º–æ–º–µ–Ω—Ç–æ–≤")
            
            # –ü–æ–ª—É—á–∞–µ–º —Å–æ–¥–µ—Ä–∂–∏–º–æ–µ –¥–æ–∫—É–º–µ–Ω—Ç–∞
            content = document.get_content_text()
            
            if not content or len(content.strip()) < 50:
                raise ValueError("–ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ —Å–æ–¥–µ—Ä–∂–∏–º–æ–≥–æ –¥–ª—è –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –∫–ª—é—á–µ–≤—ã—Ö –º–æ–º–µ–Ω—Ç–æ–≤")
            
            # –ü–æ–ª—É—á–∞–µ–º –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ —Å—Ç—Ä–æ–∫ –≤ —Ç–∞–±–ª–∏—Ü–∞—Ö –¥–æ–∫—É–º–µ–Ω—Ç–∞
            table_rows_count = self._get_table_rows_count(document)
            
            logger.info(f"–ù–∞—á–∏–Ω–∞–µ–º –≥–µ–Ω–µ—Ä–∞—Ü–∏—é –∫–ª—é—á–µ–≤—ã—Ö –º–æ–º–µ–Ω—Ç–æ–≤ –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.id} —Å —Å–æ–¥–µ—Ä–∂–∏–º—ã–º {len(content)} —Å–∏–º–≤–æ–ª–æ–≤, —Ç–∞–±–ª–∏—Ü: {table_rows_count} —Å—Ç—Ä–æ–∫")
            
            # –ì–µ–Ω–µ—Ä–∏—Ä—É–µ–º –∫–ª—é—á–µ–≤—ã–µ –º–æ–º–µ–Ω—Ç—ã —á–µ—Ä–µ–∑ Ollama —Å —É—á–µ—Ç–æ–º –∫–æ–ª–∏—á–µ—Å—Ç–≤–∞ —Å—Ç—Ä–æ–∫ —Ç–∞–±–ª–∏—Ü
            result = self.ollama_service.extract_key_points(content, table_rows_count)
            
            logger.info(f"–†–µ–∑—É–ª—å—Ç–∞—Ç –æ—Ç Ollama –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.id}: {result.get('success', False)}")
            
            if result.get("success", False):
                key_points_result = result.get("key_points_result", {})
                
                # –ò–∑–≤–ª–µ–∫–∞–µ–º –∫–ª—é—á–µ–≤—ã–µ –º–æ–º–µ–Ω—Ç—ã
                key_points = key_points_result.get("key_points", [])
                summary = key_points_result.get("summary", "")
                main_topics = key_points_result.get("main_topics", [])
                
                logger.info(f"–ü–æ–ª—É—á–µ–Ω–æ –∫–ª—é—á–µ–≤—ã—Ö –º–æ–º–µ–Ω—Ç–æ–≤: {len(key_points)}")
                
                # –°–æ—Ö—Ä–∞–Ω—è–µ–º –∫–ª—é—á–µ–≤—ã–µ –º–æ–º–µ–Ω—Ç—ã –≤ –¥–æ–∫—É–º–µ–Ω—Ç
                document.key_points = key_points
                document.key_points_generated_date = self.timezone.now()
                document.save(update_fields=['key_points', 'key_points_generated_date'])
                
                logger.info(f"–ö–ª—é—á–µ–≤—ã–µ –º–æ–º–µ–Ω—Ç—ã —Å–æ—Ö—Ä–∞–Ω–µ–Ω—ã –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.id}")
                
                return {
                    "success": True,
                    "key_points": key_points,
                    "summary": summary,
                    "main_topics": main_topics
                }
            else:
                error_msg = result.get("error", "–ù–µ–∏–∑–≤–µ—Å—Ç–Ω–∞—è –æ—à–∏–±–∫–∞ –ø—Ä–∏ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏")
                logger.error(f"Ollama –≤–µ—Ä–Ω—É–ª –æ—à–∏–±–∫—É –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.id}: {error_msg}")
                raise Exception(error_msg)
                
        except Exception as e:
            logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –∫–ª—é—á–µ–≤—ã—Ö –º–æ–º–µ–Ω—Ç–æ–≤ –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.id}: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }


class DocumentTableAnalysisService:
    """
    –°–µ—Ä–≤–∏—Å –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞ —Ç–∞–±–ª–∏—Ü –≤ –¥–æ–∫—É–º–µ–Ω—Ç–∞—Ö
    """
    
    def __init__(self):
        import re
        from django.utils import timezone
        
        self.re = re
        self.timezone = timezone
        self.logger = logging.getLogger(__name__)
    
    def analyze_document_tables(self, document: Document) -> Dict[str, Any]:
        """
        –ê–Ω–∞–ª–∏–∑–∏—Ä—É–µ—Ç –≤—Å–µ —Ç–∞–±–ª–∏—Ü—ã –≤ –¥–æ–∫—É–º–µ–Ω—Ç–µ
        
        Args:
            document: –û–±—ä–µ–∫—Ç Document
            
        Returns:
            Dict —Å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞–º–∏ –∞–Ω–∞–ª–∏–∑–∞
        """
        try:
            if not document.has_content():
                raise ValueError("–î–æ–∫—É–º–µ–Ω—Ç –¥–æ–ª–∂–µ–Ω –±—ã—Ç—å –æ–±—Ä–∞–±–æ—Ç–∞–Ω –ø–µ—Ä–µ–¥ –∞–Ω–∞–ª–∏–∑–æ–º —Ç–∞–±–ª–∏—Ü")
            
            # –ü–æ–ª—É—á–∞–µ–º –≤—Å–µ —Ç–∞–±–ª–∏—Ü—ã –¥–æ–∫—É–º–µ–Ω—Ç–∞
            tables = DocumentTable.objects.filter(document=document).order_by('order')
            
            if not tables.exists():
                return {
                    "success": True,
                    "message": "–í –¥–æ–∫—É–º–µ–Ω—Ç–µ –Ω–µ—Ç —Ç–∞–±–ª–∏—Ü –¥–ª—è –∞–Ω–∞–ª–∏–∑–∞",
                    "tables_count": 0,
                    "analyses": []
                }
            
            analyses = []
            total_metrics = {
                'total_rows': 0,
                'total_columns': 0,
                'total_cells': 0,
                'total_empty_cells': 0,
                'total_numeric_cells': 0,
                'total_text_cells': 0
            }
            
            for table in tables:
                analysis = self._analyze_single_table(table)
                analyses.append(analysis)
                
                # –°—É–º–º–∏—Ä—É–µ–º –º–µ—Ç—Ä–∏–∫–∏
                total_metrics['total_rows'] += analysis['row_count']
                total_metrics['total_columns'] += analysis['column_count']
                total_metrics['total_cells'] += analysis['cell_count']
                total_metrics['total_empty_cells'] += analysis['empty_cells_count']
                total_metrics['total_numeric_cells'] += analysis['numeric_cells_count']
                total_metrics['total_text_cells'] += analysis['text_cells_count']
            
            # –°–æ—Ö—Ä–∞–Ω—è–µ–º –∞–Ω–∞–ª–∏–∑—ã –≤ –±–∞–∑—É –¥–∞–Ω–Ω—ã—Ö
            self._save_table_analyses(document, analyses)
            
            self.logger.info(f"–ê–Ω–∞–ª–∏–∑ —Ç–∞–±–ª–∏—Ü –∑–∞–≤–µ—Ä—à–µ–Ω –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.id}: {len(analyses)} —Ç–∞–±–ª–∏—Ü")
            
            return {
                "success": True,
                "tables_count": len(analyses),
                "analyses": analyses,
                "summary": self._generate_summary(total_metrics, len(analyses))
            }
            
        except Exception as e:
            self.logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –∞–Ω–∞–ª–∏–∑–µ —Ç–∞–±–ª–∏—Ü –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.id}: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _analyze_single_table(self, table: DocumentTable) -> Dict[str, Any]:
        """
        –ê–Ω–∞–ª–∏–∑–∏—Ä—É–µ—Ç –æ–¥–Ω—É —Ç–∞–±–ª–∏—Ü—É
        
        Args:
            table: –û–±—ä–µ–∫—Ç DocumentTable
            
        Returns:
            Dict —Å —Ä–µ–∑—É–ª—å—Ç–∞—Ç–∞–º–∏ –∞–Ω–∞–ª–∏–∑–∞ —Ç–∞–±–ª–∏—Ü—ã
        """
        try:
            # –ü–∞—Ä—Å–∏–º –¥–∞–Ω–Ω—ã–µ —Ç–∞–±–ª–∏—Ü—ã
            table_data = table.data
            
            if not table_data or 'rows' not in table_data:
                return {
                    "success": False,
                    "error": "–ù–µ–∫–æ—Ä—Ä–µ–∫—Ç–Ω—ã–µ –¥–∞–Ω–Ω—ã–µ —Ç–∞–±–ª–∏—Ü—ã"
                }
            
            rows = table_data['rows']
            row_count = len(rows)
            column_count = len(rows[0]) if rows else 0
            cell_count = row_count * column_count
            
            # –ê–Ω–∞–ª–∏–∑–∏—Ä—É–µ–º —Å–æ–¥–µ—Ä–∂–∏–º–æ–µ —è—á–µ–µ–∫
            empty_cells = 0
            numeric_cells = 0
            text_cells = 0
            has_headers = False
            header_row_count = 0
            
            # –ê–Ω–∞–ª–∏–∑–∏—Ä—É–µ–º –ø–µ—Ä–≤—É—é —Å—Ç—Ä–æ–∫—É –Ω–∞ –ø—Ä–µ–¥–º–µ—Ç –∑–∞–≥–æ–ª–æ–≤–∫–æ–≤
            if rows:
                first_row = rows[0]
                header_indicators = 0
                
                for cell in first_row:
                    cell_text = str(cell).strip().lower()
                    # –ü—Ä–æ–≤–µ—Ä—è–µ–º, —è–≤–ª—è–µ—Ç—Å—è –ª–∏ —è—á–µ–π–∫–∞ –∑–∞–≥–æ–ª–æ–≤–∫–æ–º
                    if (self._is_header_cell(cell_text) or 
                        cell_text in ['‚Ññ', '–Ω–æ–º–µ—Ä', '–Ω–∞–∑–≤–∞–Ω–∏–µ', '–Ω–∞–∏–º–µ–Ω–æ–≤–∞–Ω–∏–µ', '–∫–æ–ª–∏—á–µ—Å—Ç–≤–æ', '—Å—É–º–º–∞', '–¥–∞—Ç–∞']):
                        header_indicators += 1
                
                # –ï—Å–ª–∏ –±–æ–ª—å—à–µ –ø–æ–ª–æ–≤–∏–Ω—ã —è—á–µ–µ–∫ –ø–µ—Ä–≤–æ–π —Å—Ç—Ä–æ–∫–∏ –ø–æ—Ö–æ–∂–∏ –Ω–∞ –∑–∞–≥–æ–ª–æ–≤–∫–∏
                if header_indicators > column_count / 2:
                    has_headers = True
                    header_row_count = 1
            
            # –ê–Ω–∞–ª–∏–∑–∏—Ä—É–µ–º –≤—Å–µ —è—á–µ–π–∫–∏
            for row in rows:
                for cell in row:
                    cell_text = str(cell).strip()
                    
                    if not cell_text:
                        empty_cells += 1
                    elif self._is_numeric(cell_text):
                        numeric_cells += 1
                    else:
                        text_cells += 1
            
            # –û–ø—Ä–µ–¥–µ–ª—è–µ–º —Ç–∏–ø —Ç–∞–±–ª–∏—Ü—ã
            table_type = self._determine_table_type(rows, numeric_cells, text_cells)
            
            # –û–ø—Ä–µ–¥–µ–ª—è–µ–º –æ—Å–Ω–æ–≤–Ω—É—é —Ç–µ–º—É
            main_topic = self._extract_main_topic(rows)
            
            # –ò–∑–≤–ª–µ–∫–∞–µ–º –∫–ª—é—á–µ–≤—ã–µ –º–µ—Ç—Ä–∏–∫–∏
            key_metrics = self._extract_key_metrics(rows, numeric_cells)
            
            return {
                "success": True,
                "table_id": table.id,
                "table_title": table.title,
                "row_count": row_count,
                "column_count": column_count,
                "cell_count": cell_count,
                "empty_cells_count": empty_cells,
                "numeric_cells_count": numeric_cells,
                "text_cells_count": text_cells,
                "has_headers": has_headers,
                "header_row_count": header_row_count,
                "table_type": table_type,
                "main_topic": main_topic,
                "key_metrics": key_metrics,
                "fill_percentage": round(((cell_count - empty_cells) / cell_count) * 100, 2) if cell_count > 0 else 0,
                "numeric_percentage": round((numeric_cells / cell_count) * 100, 2) if cell_count > 0 else 0,
                "text_percentage": round((text_cells / cell_count) * 100, 2) if cell_count > 0 else 0
            }
            
        except Exception as e:
            self.logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –∞–Ω–∞–ª–∏–∑–µ —Ç–∞–±–ª–∏—Ü—ã {table.id}: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _is_numeric(self, text: str) -> bool:
        """–ü—Ä–æ–≤–µ—Ä—è–µ—Ç, —è–≤–ª—è–µ—Ç—Å—è –ª–∏ —Ç–µ–∫—Å—Ç —á–∏—Å–ª–æ–≤—ã–º"""
        try:
            # –£–±–∏—Ä–∞–µ–º –ø—Ä–æ–±–µ–ª—ã –∏ –∑–∞–º–µ–Ω—è–µ–º –∑–∞–ø—è—Ç—ã–µ –Ω–∞ —Ç–æ—á–∫–∏
            cleaned_text = text.replace(',', '.').replace(' ', '')
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º –Ω–∞ —á–∏—Å–ª–æ
            float(cleaned_text)
            return True
        except (ValueError, TypeError):
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º –Ω–∞ –ø—Ä–æ—Ü–µ–Ω—Ç
            if cleaned_text.endswith('%'):
                try:
                    float(cleaned_text[:-1])
                    return True
                except ValueError:
                    pass
            
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º –Ω–∞ –≤–∞–ª—é—Ç—É (—Ä—É–±–ª–∏, –¥–æ–ª–ª–∞—Ä—ã –∏ —Ç.–¥.)
            currency_patterns = [
                r'^\d+[\s,]*\d*[\s,]*\d*\s*(—Ä—É–±|—Ä\.|‚ÇΩ|–¥–æ–ª|usd|\$|‚Ç¨|eur)',
                r'^\d+[\s,]*\d*[\s,]*\d*\s*(—Ç—ã—Å|–º–ª–Ω|–º–ª—Ä–¥)',
                r'^\d+[\s,]*\d*[\s,]*\d*\s*(—Ç—ã—Å\.|–º–ª–Ω\.|–º–ª—Ä–¥\.)'
            ]
            
            for pattern in currency_patterns:
                if self.re.match(pattern, cleaned_text.lower()):
                    return True
            
            return False
    
    def _is_header_cell(self, text: str) -> bool:
        """–ü—Ä–æ–≤–µ—Ä—è–µ—Ç, —è–≤–ª—è–µ—Ç—Å—è –ª–∏ —è—á–µ–π–∫–∞ –∑–∞–≥–æ–ª–æ–≤–∫–æ–º"""
        header_keywords = [
            '–Ω–∞–∑–≤–∞–Ω–∏–µ', '–Ω–∞–∏–º–µ–Ω–æ–≤–∞–Ω–∏–µ', '–∏–º—è', '–∑–∞–≥–æ–ª–æ–≤–æ–∫',
            '–Ω–æ–º–µ—Ä', '‚Ññ', '–∫–æ–¥', 'id',
            '–¥–∞—Ç–∞', '–≤—Ä–µ–º—è', '–ø–µ—Ä–∏–æ–¥',
            '–∫–æ–ª–∏—á–µ—Å—Ç–≤–æ', '—Å—É–º–º–∞', '—Å—Ç–æ–∏–º–æ—Å—Ç—å', '—Ü–µ–Ω–∞',
            '—Å—Ç–∞—Ç—É—Å', '—Å–æ—Å—Ç–æ—è–Ω–∏–µ', '—Ä–µ–∑—É–ª—å—Ç–∞—Ç',
            '–æ–ø–∏—Å–∞–Ω–∏–µ', '–∫–æ–º–º–µ–Ω—Ç–∞—Ä–∏–π', '–ø—Ä–∏–º–µ—á–∞–Ω–∏–µ'
        ]
        
        text_lower = text.lower()
        return any(keyword in text_lower for keyword in header_keywords)
    
    def _determine_table_type(self, rows: List[List], numeric_cells: int, text_cells: int) -> str:
        """–û–ø—Ä–µ–¥–µ–ª—è–µ—Ç —Ç–∏–ø —Ç–∞–±–ª–∏—Ü—ã"""
        if not rows:
            return "–ø—É—Å—Ç–∞—è"
        
        total_cells = len(rows) * len(rows[0]) if rows else 0
        
        if total_cells == 0:
            return "–ø—É—Å—Ç–∞—è"
        
        numeric_ratio = numeric_cells / total_cells
        
        if numeric_ratio > 0.7:
            return "—á–∏—Å–ª–æ–≤–∞—è"
        elif numeric_ratio > 0.3:
            return "—Å–º–µ—à–∞–Ω–Ω–∞—è"
        else:
            return "—Ç–µ–∫—Å—Ç–æ–≤–∞—è"
    
    def _extract_main_topic(self, rows: List[List]) -> str:
        """–ò–∑–≤–ª–µ–∫–∞–µ—Ç –æ—Å–Ω–æ–≤–Ω—É—é —Ç–µ–º—É —Ç–∞–±–ª–∏—Ü—ã"""
        if not rows or not rows[0]:
            return "–Ω–µ–∏–∑–≤–µ—Å—Ç–Ω–æ"
        
        # –ë–µ—Ä–µ–º –ø–µ—Ä–≤—É—é —Å—Ç—Ä–æ–∫—É –∫–∞–∫ –ø–æ—Ç–µ–Ω—Ü–∏–∞–ª—å–Ω—ã–µ –∑–∞–≥–æ–ª–æ–≤–∫–∏
        first_row = rows[0]
        topics = []
        
        for cell in first_row:
            cell_text = str(cell).strip()
            if cell_text and len(cell_text) > 3:
                topics.append(cell_text)
        
        if topics:
            # –í–æ–∑–≤—Ä–∞—â–∞–µ–º —Å–∞–º—ã–π –¥–ª–∏–Ω–Ω—ã–π –∑–∞–≥–æ–ª–æ–≤–æ–∫ –∫–∞–∫ –æ—Å–Ω–æ–≤–Ω—É—é —Ç–µ–º—É
            return max(topics, key=len)
        
        return "–Ω–µ–∏–∑–≤–µ—Å—Ç–Ω–æ"
    
    def _extract_key_metrics(self, rows: List[List], numeric_cells: int) -> List[Dict[str, Any]]:
        """–ò–∑–≤–ª–µ–∫–∞–µ—Ç –∫–ª—é—á–µ–≤—ã–µ –º–µ—Ç—Ä–∏–∫–∏ –∏–∑ —Ç–∞–±–ª–∏—Ü—ã"""
        metrics = []
        
        if not rows:
            return metrics
        
        # –ë–∞–∑–æ–≤–∞—è —Å—Ç–∞—Ç–∏—Å—Ç–∏–∫–∞
        metrics.append({
            "name": "–û–±—â–µ–µ –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ —Å—Ç—Ä–æ–∫",
            "value": len(rows),
            "type": "count"
        })
        
        metrics.append({
            "name": "–ö–æ–ª–∏—á–µ—Å—Ç–≤–æ —Å—Ç–æ–ª–±—Ü–æ–≤",
            "value": len(rows[0]) if rows else 0,
            "type": "count"
        })
        
        metrics.append({
            "name": "–ß–∏—Å–ª–æ–≤—ã—Ö —è—á–µ–µ–∫",
            "value": numeric_cells,
            "type": "count"
        })
        
        # –ï—Å–ª–∏ –µ—Å—Ç—å —á–∏—Å–ª–æ–≤—ã–µ –¥–∞–Ω–Ω—ã–µ, –ø—ã—Ç–∞–µ–º—Å—è –Ω–∞–π—Ç–∏ —Å—É–º–º—ã –∏ —Å—Ä–µ–¥–Ω–∏–µ
        if numeric_cells > 0:
            numeric_values = []
            for row in rows:
                for cell in row:
                    cell_text = str(cell).strip()
                    if self._is_numeric(cell_text):
                        try:
                            cleaned = cell_text.replace(',', '.').replace(' ', '')
                            # –£–±–∏—Ä–∞–µ–º –≤–∞–ª—é—Ç–Ω—ã–µ —Å–∏–º–≤–æ–ª—ã
                            cleaned = self.re.sub(r'[^\d.,]', '', cleaned)
                            if cleaned:
                                numeric_values.append(float(cleaned))
                        except ValueError:
                            continue
            
            if numeric_values:
                metrics.append({
                    "name": "–°—É–º–º–∞ –≤—Å–µ—Ö –∑–Ω–∞—á–µ–Ω–∏–π",
                    "value": sum(numeric_values),
                    "type": "sum"
                })
                
                metrics.append({
                    "name": "–°—Ä–µ–¥–Ω–µ–µ –∑–Ω–∞—á–µ–Ω–∏–µ",
                    "value": sum(numeric_values) / len(numeric_values),
                    "type": "average"
                })
                
                metrics.append({
                    "name": "–ú–∞–∫—Å–∏–º–∞–ª—å–Ω–æ–µ –∑–Ω–∞—á–µ–Ω–∏–µ",
                    "value": max(numeric_values),
                    "type": "max"
                })
                
                metrics.append({
                    "name": "–ú–∏–Ω–∏–º–∞–ª—å–Ω–æ–µ –∑–Ω–∞—á–µ–Ω–∏–µ",
                    "value": min(numeric_values),
                    "type": "min"
                })
        
        return metrics
    
    def _generate_summary(self, total_metrics: Dict[str, int], tables_count: int) -> Dict[str, Any]:
        """–ì–µ–Ω–µ—Ä–∏—Ä—É–µ—Ç –æ–±—â—É—é —Å–≤–æ–¥–∫—É –ø–æ –≤—Å–µ–º —Ç–∞–±–ª–∏—Ü–∞–º"""
        if tables_count == 0:
            return {
                "message": "–í –¥–æ–∫—É–º–µ–Ω—Ç–µ –Ω–µ—Ç —Ç–∞–±–ª–∏—Ü",
                "tables_count": 0
            }
        
        avg_rows = total_metrics['total_rows'] / tables_count if tables_count > 0 else 0
        avg_columns = total_metrics['total_columns'] / tables_count if tables_count > 0 else 0
        
        fill_percentage = 0
        if total_metrics['total_cells'] > 0:
            filled_cells = total_metrics['total_cells'] - total_metrics['total_empty_cells']
            fill_percentage = round((filled_cells / total_metrics['total_cells']) * 100, 2)
        
        return {
            "tables_count": tables_count,
            "total_rows": total_metrics['total_rows'],
            "total_columns": total_metrics['total_columns'],
            "total_cells": total_metrics['total_cells'],
            "avg_rows_per_table": round(avg_rows, 1),
            "avg_columns_per_table": round(avg_columns, 1),
            "fill_percentage": fill_percentage,
            "numeric_percentage": round((total_metrics['total_numeric_cells'] / total_metrics['total_cells']) * 100, 2) if total_metrics['total_cells'] > 0 else 0,
            "text_percentage": round((total_metrics['total_text_cells'] / total_metrics['total_cells']) * 100, 2) if total_metrics['total_cells'] > 0 else 0
        }
    
    def _save_table_analyses(self, document: Document, analyses: List[Dict[str, Any]]) -> None:
        """–°–æ—Ö—Ä–∞–Ω—è–µ—Ç –∞–Ω–∞–ª–∏–∑—ã —Ç–∞–±–ª–∏—Ü –≤ –±–∞–∑—É –¥–∞–Ω–Ω—ã—Ö"""
        try:
            # –£–¥–∞–ª—è–µ–º —Å—Ç–∞—Ä—ã–µ –∞–Ω–∞–ª–∏–∑—ã
            DocumentTableAnalysis.objects.filter(document=document).delete()
            
            # –°–æ—Ö—Ä–∞–Ω—è–µ–º –Ω–æ–≤—ã–µ –∞–Ω–∞–ª–∏–∑—ã
            for analysis in analyses:
                if analysis.get('success', False):
                    table = DocumentTable.objects.get(id=analysis['table_id']) if analysis.get('table_id') else None
                    
                    DocumentTableAnalysis.objects.create(
                        document=document,
                        table=table,
                        row_count=analysis['row_count'],
                        column_count=analysis['column_count'],
                        cell_count=analysis['cell_count'],
                        empty_cells_count=analysis['empty_cells_count'],
                        numeric_cells_count=analysis['numeric_cells_count'],
                        text_cells_count=analysis['text_cells_count'],
                        has_headers=analysis['has_headers'],
                        header_row_count=analysis['header_row_count'],
                        table_type=analysis['table_type'],
                        main_topic=analysis['main_topic'],
                        key_metrics=analysis['key_metrics']
                    )
            
            self.logger.info(f"–ê–Ω–∞–ª–∏–∑—ã —Ç–∞–±–ª–∏—Ü —Å–æ—Ö—Ä–∞–Ω–µ–Ω—ã –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.id}")
            
        except Exception as e:
            self.logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ —Å–æ—Ö—Ä–∞–Ω–µ–Ω–∏–∏ –∞–Ω–∞–ª–∏–∑–æ–≤ —Ç–∞–±–ª–∏—Ü –¥–ª—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ {document.id}: {str(e)}")
            raise
