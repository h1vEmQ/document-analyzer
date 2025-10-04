#!/usr/bin/env python
"""
Скрипт для тестирования системы анализа изменений
"""
import os
import sys
import django

# Настройка Django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'wara_project.settings')
django.setup()

from analysis.services import DocumentComparisonService, AnalysisSettingsService
from analysis.models import Comparison
from documents.models import Document as DjangoDocument
from users.models import User
from django.utils import timezone

def test_document_analysis():
    """Тестирование системы анализа документов"""
    
    print("🔍 Тестирование системы анализа изменений WARA")
    print("=" * 60)
    
    # Получаем пользователя
    user = User.objects.first()
    if not user:
        print("❌ Пользователь не найден!")
        return
    
    # Получаем обработанные документы
    processed_docs = DjangoDocument.objects.filter(
        user=user,
        status='processed'
    ).order_by('id')
    
    if len(processed_docs) < 2:
        print("❌ Недостаточно обработанных документов для сравнения!")
        print(f"   Найдено документов: {len(processed_docs)}")
        print("   Требуется минимум 2 документа")
        return
    
    base_doc = processed_docs[0]
    compared_doc = processed_docs[1]
    
    print(f"📄 Базовый документ: {base_doc.title}")
    print(f"📄 Сравниваемый документ: {compared_doc.title}")
    
    # Создаем сравнение
    print("\n💾 Создание сравнения в БД...")
    
    comparison, created = Comparison.objects.get_or_create(
        user=user,
        base_document=base_doc,
        compared_document=compared_doc,
        defaults={
            'status': 'pending'
        }
    )
    
    if created:
        print(f"✅ Сравнение создано с ID: {comparison.id}")
    else:
        print(f"✅ Используется существующее сравнение ID: {comparison.id}")
    
    # Тестируем анализ
    print("\n🔧 Тестирование анализа изменений...")
    
    try:
        comparison_service = DocumentComparisonService()
        
        # Выполняем анализ
        start_time = timezone.now()
        analysis_result = comparison_service.compare_documents(comparison)
        end_time = timezone.now()
        
        processing_time = (end_time - start_time).total_seconds()
        
        print("✅ Анализ успешно завершен!")
        print(f"⏱️ Время обработки: {processing_time:.2f} секунд")
        
        # Выводим результаты
        print(f"\n📊 Результаты анализа:")
        
        summary = analysis_result.get('summary', {})
        print(f"   • Добавлено: {summary.get('added', 0)}")
        print(f"   • Удалено: {summary.get('removed', 0)}")
        print(f"   • Изменено: {summary.get('modified', 0)}")
        print(f"   • Всего изменений: {summary.get('total', 0)}")
        
        # Детальные результаты
        print(f"\n🔍 Детальные изменения:")
        
        for change_type, changes in analysis_result.items():
            if change_type.endswith('_changes') and isinstance(changes, list):
                if changes:
                    print(f"\n📋 {change_type.replace('_', ' ').title()}:")
                    for i, change in enumerate(changes[:3], 1):  # Показываем первые 3
                        print(f"   {i}. {change.get('type', 'unknown')} в {change.get('section', 'unknown')}")
                        print(f"      {change.get('content', '')[:100]}...")
                    
                    if len(changes) > 3:
                        print(f"   ... и еще {len(changes) - 3} изменений")
        
        # Сохраняем результаты в БД
        print(f"\n💾 Сохранение результатов в БД...")
        comparison_service.save_comparison_results(comparison, analysis_result)
        
        # Проверяем сохранение
        comparison.refresh_from_db()
        print(f"✅ Статус сравнения: {comparison.status}")
        print(f"✅ Количество изменений в БД: {comparison.changes.count()}")
        
        # Показываем статистику изменений
        if comparison.changes.exists():
            print(f"\n📈 Статистика изменений:")
            changes_by_type = comparison.changes.values('change_type').distinct()
            for change_type in changes_by_type:
                count = comparison.changes.filter(change_type=change_type['change_type']).count()
                print(f"   • {change_type['change_type']}: {count}")
        
        # Тестируем настройки анализа
        print(f"\n⚙️ Тестирование настроек анализа...")
        settings_service = AnalysisSettingsService()
        user_settings = settings_service.get_user_settings(user)
        
        print(f"✅ Настройки пользователя:")
        print(f"   • Чувствительность: {user_settings.sensitivity}")
        print(f"   • Анализировать текстовые изменения: {user_settings.include_text_changes}")
        print(f"   • Анализировать изменения таблиц: {user_settings.include_table_changes}")
        print(f"   • Анализировать структурные изменения: {user_settings.include_structure_changes}")
        print(f"   • Минимальная длина изменения: {user_settings.min_change_length}")
        
    except Exception as e:
        print(f"❌ Ошибка при анализе: {str(e)}")
        import traceback
        traceback.print_exc()
    
    print(f"\n🎉 Тестирование анализа завершено!")
    print(f"📄 Сравнение ID: {comparison.id}")

def create_test_documents():
    """Создание тестовых документов для сравнения"""
    
    print("📝 Создание тестовых документов...")
    
    from docx import Document
    from django.core.files import File
    import tempfile
    
    user = User.objects.first()
    if not user:
        print("❌ Пользователь не найден!")
        return
    
    # Создаем первый документ
    doc1 = Document()
    doc1.add_heading('Отчет за неделю 1', 0)
    doc1.add_heading('1. Выполненные задачи', level=1)
    doc1.add_paragraph('Задача 1: Разработка модуля парсинга')
    doc1.add_paragraph('Задача 2: Создание интерфейса')
    doc1.add_heading('2. Результаты', level=1)
    doc1.add_paragraph('Все задачи выполнены в срок')
    
    # Создаем второй документ (с изменениями)
    doc2 = Document()
    doc2.add_heading('Отчет за неделю 2', 0)
    doc2.add_heading('1. Выполненные задачи', level=1)
    doc2.add_paragraph('Задача 1: Разработка модуля парсинга')
    doc2.add_paragraph('Задача 2: Создание интерфейса')
    doc2.add_paragraph('Задача 3: Тестирование системы')  # Добавлено
    doc2.add_heading('2. Результаты', level=1)
    doc2.add_paragraph('Все задачи выполнены в срок')
    doc2.add_heading('3. Планы на следующую неделю', level=1)  # Добавлен новый раздел
    doc2.add_paragraph('Планируется оптимизация производительности')
    
    # Сохраняем документы
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f1:
        doc1.save(f1.name)
        with open(f1.name, 'rb') as file1:
            django_file1 = File(file1, name='test_week1.docx')
            document1 = DjangoDocument(
                title='Тестовый отчет - неделя 1',
                filename='test_week1.docx',
                file=django_file1,
                file_size=os.path.getsize(f1.name),
                user=user
            )
            document1.save()
            print(f"✅ Создан документ 1: {document1.id}")
    
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f2:
        doc2.save(f2.name)
        with open(f2.name, 'rb') as file2:
            django_file2 = File(file2, name='test_week2.docx')
            document2 = DjangoDocument(
                title='Тестовый отчет - неделя 2',
                filename='test_week2.docx',
                file=django_file2,
                file_size=os.path.getsize(f2.name),
                user=user
            )
            document2.save()
            print(f"✅ Создан документ 2: {document2.id}")
    
    # Обрабатываем документы парсером
    from documents.services import DocumentParserService
    
    parser_service = DocumentParserService()
    
    print("🔧 Обработка документов парсером...")
    
    try:
        # Обрабатываем первый документ
        content_data1 = parser_service.parse_document(document1)
        parser_service.save_parsed_content(document1, content_data1)
        print(f"✅ Документ 1 обработан")
        
        # Обрабатываем второй документ
        content_data2 = parser_service.parse_document(document2)
        parser_service.save_parsed_content(document2, content_data2)
        print(f"✅ Документ 2 обработан")
        
    except Exception as e:
        print(f"❌ Ошибка при обработке документов: {str(e)}")

if __name__ == '__main__':
    # Сначала создаем тестовые документы, если их нет
    processed_docs = DjangoDocument.objects.filter(status='processed')
    if len(processed_docs) < 2:
        print("📝 Создание тестовых документов...")
        create_test_documents()
        print()
    
    # Затем тестируем анализ
    test_document_analysis()
