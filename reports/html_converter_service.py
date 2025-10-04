"""
Сервис для конвертации отчетов в HTML для просмотра в браузере
"""
import os
import logging
from typing import Optional, Dict, Any
from django.conf import settings
from django.core.files.base import ContentFile
from django.utils.html import escape
import re
import PyPDF2
from docx import Document
import io

logger = logging.getLogger(__name__)


class HTMLReportConverterService:
    """
    Сервис для конвертации отчетов в HTML формат
    """
    
    def __init__(self):
        self.base_styles = self._get_base_styles()
    
    def _get_base_styles(self) -> str:
        """Возвращает базовые CSS стили для HTML отчета"""
        return """
        <style>
            .report-container {
                max-width: 1200px;
                margin: 0 auto;
                padding: 10px;
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                line-height: 1.4;
                color: #333;
                font-size: 14px;
            }
            .report-header {
                border-bottom: 1px solid #007bff;
                padding-bottom: 10px;
                margin-bottom: 15px;
            }
            .report-title {
                color: #007bff;
                font-size: 1.8em;
                font-weight: bold;
                margin-bottom: 5px;
            }
            .report-meta {
                display: flex;
                gap: 20px;
                flex-wrap: wrap;
                margin-bottom: 10px;
            }
            .meta-item {
                display: flex;
                flex-direction: column;
            }
            .meta-label {
                font-weight: bold;
                color: #666;
                font-size: 0.8em;
                text-transform: uppercase;
                letter-spacing: 0.3px;
            }
            .meta-value {
                color: #333;
                font-size: 0.95em;
            }
            .report-section {
                margin-bottom: 15px;
                padding: 10px;
                border: 1px solid #e9ecef;
                border-radius: 4px;
                background: #f8f9fa;
            }
            .section-title {
                color: #007bff;
                font-size: 1.4em;
                font-weight: bold;
                margin-bottom: 8px;
                border-bottom: 1px solid #dee2e6;
                padding-bottom: 5px;
            }
            .subsection-title {
                color: #495057;
                font-size: 1.1em;
                font-weight: bold;
                margin-bottom: 6px;
                margin-top: 10px;
            }
            .content-text {
                margin-bottom: 8px;
                text-align: justify;
                font-size: 0.95em;
            }
            .highlight-box {
                background: #fff3cd;
                border: 1px solid #ffeaa7;
                border-radius: 3px;
                padding: 8px;
                margin: 8px 0;
                font-size: 0.9em;
            }
            .info-table {
                width: 100%;
                border-collapse: collapse;
                margin: 8px 0;
                background: white;
                border-radius: 3px;
                overflow: hidden;
                box-shadow: 0 1px 2px rgba(0,0,0,0.1);
                font-size: 0.9em;
            }
            .info-table th,
            .info-table td {
                padding: 6px 8px;
                text-align: left;
                border-bottom: 1px solid #dee2e6;
            }
            .info-table th {
                background: #007bff;
                color: white;
                font-weight: bold;
                font-size: 0.85em;
            }
            .info-table tr:hover {
                background: #f8f9fa;
            }
            .list-item {
                margin-bottom: 4px;
                padding-left: 15px;
                position: relative;
                font-size: 0.9em;
            }
            .list-item::before {
                content: "•";
                color: #007bff;
                font-weight: bold;
                position: absolute;
                left: 0;
            }
            .badge {
                display: inline-block;
                padding: 2px 6px;
                font-size: 0.7em;
                font-weight: bold;
                border-radius: 3px;
                margin-right: 3px;
            }
            .badge-success { background: #d4edda; color: #155724; }
            .badge-warning { background: #fff3cd; color: #856404; }
            .badge-danger { background: #f8d7da; color: #721c24; }
            .badge-info { background: #d1ecf1; color: #0c5460; }
            .badge-secondary { background: #e2e3e5; color: #383d41; }
            .document-comparison {
                display: grid;
                grid-template-columns: 1fr 1fr;
                gap: 10px;
                margin: 10px 0;
            }
            .document-section {
                padding: 8px;
                border: 1px solid #dee2e6;
                border-radius: 3px;
                background: white;
                font-size: 0.9em;
            }
            .document-title {
                font-weight: bold;
                color: #007bff;
                margin-bottom: 5px;
                font-size: 1em;
            }
            h3, h4 {
                margin: 8px 0 4px 0;
                font-size: 1.1em;
            }
            p {
                margin: 4px 0;
                font-size: 0.95em;
            }
            @media (max-width: 768px) {
                .document-comparison {
                    grid-template-columns: 1fr;
                }
                .report-meta {
                    flex-direction: column;
                    gap: 8px;
                }
                .report-container {
                    padding: 5px;
                }
            }
        </style>
        """
    
    def convert_report_to_html(self, report) -> str:
        """
        Конвертирует отчет в HTML для просмотра в браузере
        
        Args:
            report: Объект отчета
            
        Returns:
            str: HTML содержимое отчета
        """
        try:
            # Получаем данные отчета в зависимости от формата
            if report.format.lower() == 'pdf':
                return self._convert_pdf_to_html(report)
            elif report.format.lower() == 'docx':
                return self._convert_docx_to_html(report)
            else:
                return self._create_fallback_html(report)
                
        except Exception as e:
            logger.error(f"Ошибка конвертации отчета {report.id} в HTML: {e}")
            return self._create_error_html(report, str(e))
    
    def _extract_pdf_content(self, report) -> str:
        """Извлекает текстовое содержимое из PDF файла"""
        try:
            if not report.file or not report.file.path:
                return "Файл отчета не найден"
            
            with open(report.file.path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = ""
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            # Убираем заголовок страницы для компактности, если это не первая страница
                            if page_num > 0:
                                text_content += f"<div class='content-text'>{self._format_text_content(page_text)}</div>\n"
                            else:
                                text_content += f"<h4 class='subsection-title'>Страница {page_num + 1}</h4>\n"
                                text_content += f"<div class='content-text'>{self._format_text_content(page_text)}</div>\n"
                    except Exception as e:
                        logger.warning(f"Ошибка извлечения текста со страницы {page_num + 1}: {e}")
                        text_content += f"<p class='text-muted'>Не удалось извлечь текст со страницы {page_num + 1}</p>\n"
                
                return text_content if text_content else "Не удалось извлечь текстовое содержимое из PDF"
                
        except Exception as e:
            logger.error(f"Ошибка извлечения содержимого из PDF {report.id}: {e}")
            return f"Ошибка чтения PDF файла: {str(e)}"
    
    def _extract_docx_content(self, report) -> str:
        """Извлекает текстовое содержимое из DOCX файла"""
        try:
            if not report.file or not report.file.path:
                return "Файл отчета не найден"
            
            doc = Document(report.file.path)
            content = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # Проверяем стиль параграфа для определения типа контента
                    if (paragraph.style.name.startswith('Heading') or 
                        (len(paragraph.text.strip()) < 100 and paragraph.text.strip().isupper())):
                        content += f"<h4 class='subsection-title'>{self._format_text_content(paragraph.text)}</h4>\n"
                    else:
                        content += f"<p class='content-text'>{self._format_text_content(paragraph.text)}</p>\n"
            
            # Извлекаем содержимое таблиц
            for table in doc.tables:
                content += "<table class='info-table'>\n"
                for i, row in enumerate(table.rows):
                    content += "<tr>\n"
                    for cell in row.cells:
                        tag = "th" if i == 0 else "td"
                        content += f"<{tag}>{self._format_text_content(cell.text)}</{tag}>\n"
                    content += "</tr>\n"
                content += "</table>\n"
            
            return content if content else "Документ не содержит текстового содержимого"
            
        except Exception as e:
            logger.error(f"Ошибка извлечения содержимого из DOCX {report.id}: {e}")
            return f"Ошибка чтения DOCX файла: {str(e)}"
    
    def _format_text_content(self, text: str) -> str:
        """Форматирует текстовое содержимое для HTML"""
        if not text:
            return ""
        
        # Очищаем текст от лишних пробелов и символов
        text = text.strip()
        
        # Удаляем повторяющиеся пробелы и переносы строк
        text = re.sub(r'\s+', ' ', text)
        
        # Экранируем HTML
        text = escape(text)
        
        # Разбиваем на абзацы по переносам строк
        paragraphs = text.split('\n')
        formatted_paragraphs = []
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            # Если абзац короткий и в верхнем регистре - делаем заголовком
            if (len(paragraph) < 60 and 
                paragraph.isupper() and 
                any(c.isalpha() for c in paragraph) and
                not paragraph.endswith('.') and
                not paragraph.endswith(':')):
                formatted_paragraphs.append(f"<h4 class='subsection-title'>{paragraph}</h4>")
            else:
                formatted_paragraphs.append(f"<p class='content-text'>{paragraph}</p>")
        
        return '\n'.join(formatted_paragraphs)
    
    def _convert_pdf_to_html(self, report) -> str:
        """Конвертирует PDF отчет в HTML с полным содержимым"""
        # Извлекаем содержимое из PDF
        pdf_content = self._extract_pdf_content(report)
        
        html_content = f"""
        <!DOCTYPE html>
        <html lang="ru">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{escape(report.title)}</title>
            {self.base_styles}
        </head>
        <body>
            <div class="report-container">
                <div class="report-header">
                    <h1 class="report-title">{escape(report.title)}</h1>
                    <div class="report-meta">
                        <div class="meta-item">
                            <span class="meta-label">Формат</span>
                            <span class="meta-value">PDF</span>
                        </div>
                        <div class="meta-item">
                            <span class="meta-label">Дата создания</span>
                            <span class="meta-value">{report.generated_date.strftime('%d.%m.%Y %H:%M')}</span>
                        </div>
                        <div class="meta-item">
                            <span class="meta-label">Статус</span>
                            <span class="meta-value">
                                <span class="badge badge-{'success' if report.status == 'ready' else 'warning'}">
                                    {report.get_status_display()}
                                </span>
                            </span>
                        </div>
                        <div class="meta-item">
                            <span class="meta-label">Размер файла</span>
                            <span class="meta-value">{report.get_file_size_mb()} МБ</span>
                        </div>
                    </div>
                </div>
                
                <div class="report-section">
                    <h2 class="section-title">Информация об отчете</h2>
                    <table class="info-table">
                        <tr>
                            <th>Параметр</th>
                            <th>Значение</th>
                        </tr>
                        <tr>
                            <td>Название</td>
                            <td>{escape(report.title)}</td>
                        </tr>
                        <tr>
                            <td>Формат</td>
                            <td>{report.format.upper()}</td>
                        </tr>
                        <tr>
                            <td>Шаблон</td>
                            <td>{escape(report.template_used or 'Не указан')}</td>
                        </tr>
                        <tr>
                            <td>Версия</td>
                            <td>{report.version}</td>
                        </tr>
                        <tr>
                            <td>Дата генерации</td>
                            <td>{report.generated_date.strftime('%d.%m.%Y %H:%M')}</td>
                        </tr>
                        <tr>
                            <td>Статус</td>
                            <td>
                                <span class="badge badge-{'success' if report.status == 'ready' else 'warning'}">
                                    {report.get_status_display()}
                                </span>
                            </td>
                        </tr>
                        <tr>
                            <td>Размер файла</td>
                            <td>{report.get_file_size_mb()} МБ</td>
                        </tr>
                    </table>
                </div>
                
                {self._get_comparison_info_html(report)}
                
                <div class="report-section">
                    <h2 class="section-title">Содержимое отчета</h2>
                    {pdf_content}
                </div>
                
                <div class="highlight-box">
                    <strong>Примечание:</strong> Это PDF отчет, конвертированный в HTML для просмотра в браузере.
                </div>
            </div>
        </body>
        </html>
        """
        
        return html_content
    
    def _convert_docx_to_html(self, report) -> str:
        """Конвертирует DOCX отчет в HTML с полным содержимым"""
        # Извлекаем содержимое из DOCX
        docx_content = self._extract_docx_content(report)
        
        html_content = f"""
        <!DOCTYPE html>
        <html lang="ru">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{escape(report.title)}</title>
            {self.base_styles}
        </head>
        <body>
            <div class="report-container">
                <div class="report-header">
                    <h1 class="report-title">{escape(report.title)}</h1>
                    <div class="report-meta">
                        <div class="meta-item">
                            <span class="meta-label">Формат</span>
                            <span class="meta-value">DOCX</span>
                        </div>
                        <div class="meta-item">
                            <span class="meta-label">Дата создания</span>
                            <span class="meta-value">{report.generated_date.strftime('%d.%m.%Y %H:%M')}</span>
                        </div>
                        <div class="meta-item">
                            <span class="meta-label">Статус</span>
                            <span class="meta-value">
                                <span class="badge badge-{'success' if report.status == 'ready' else 'warning'}">
                                    {report.get_status_display()}
                                </span>
                            </span>
                        </div>
                        <div class="meta-item">
                            <span class="meta-label">Размер файла</span>
                            <span class="meta-value">{report.get_file_size_mb()} МБ</span>
                        </div>
                    </div>
                </div>
                
                <div class="report-section">
                    <h2 class="section-title">Информация об отчете</h2>
                    <table class="info-table">
                        <tr>
                            <th>Параметр</th>
                            <th>Значение</th>
                        </tr>
                        <tr>
                            <td>Название</td>
                            <td>{escape(report.title)}</td>
                        </tr>
                        <tr>
                            <td>Формат</td>
                            <td>{report.format.upper()}</td>
                        </tr>
                        <tr>
                            <td>Шаблон</td>
                            <td>{escape(report.template_used or 'Не указан')}</td>
                        </tr>
                        <tr>
                            <td>Версия</td>
                            <td>{report.version}</td>
                        </tr>
                        <tr>
                            <td>Дата генерации</td>
                            <td>{report.generated_date.strftime('%d.%m.%Y %H:%M')}</td>
                        </tr>
                        <tr>
                            <td>Статус</td>
                            <td>
                                <span class="badge badge-{'success' if report.status == 'ready' else 'warning'}">
                                    {report.get_status_display()}
                                </span>
                            </td>
                        </tr>
                        <tr>
                            <td>Размер файла</td>
                            <td>{report.get_file_size_mb()} МБ</td>
                        </tr>
                    </table>
                </div>
                
                {self._get_comparison_info_html(report)}
                
                <div class="report-section">
                    <h2 class="section-title">Содержимое отчета</h2>
                    {docx_content}
                </div>
                
                <div class="highlight-box">
                    <strong>Примечание:</strong> Это DOCX отчет, конвертированный в HTML для просмотра в браузере.
                </div>
            </div>
        </body>
        </html>
        """
        
        return html_content
    
    def _get_comparison_info_html(self, report) -> str:
        """Возвращает HTML с информацией о сравнении"""
        if not report.comparison:
            return ""
        
        comparison = report.comparison
        
        html = f"""
        <div class="report-section">
            <h2 class="section-title">Информация о сравнении</h2>
            <table class="info-table">
                <tr>
                    <th>Параметр</th>
                    <th>Значение</th>
                </tr>
                <tr>
                    <td>Название сравнения</td>
                    <td>{escape(comparison.title)}</td>
                </tr>
                <tr>
                    <td>Тип анализа</td>
                    <td>{escape(comparison.get_analysis_type_display())}</td>
                </tr>
                <tr>
                    <td>Метод анализа</td>
                    <td>{escape(comparison.analysis_method or 'Не указан')}</td>
                </tr>
                <tr>
                    <td>Статус анализа</td>
                    <td>
                        <span class="badge badge-{'success' if comparison.status == 'completed' else 'warning'}">
                            {comparison.get_status_display()}
                        </span>
                    </td>
                </tr>
                <tr>
                    <td>Дата создания</td>
                    <td>{comparison.created_date.strftime('%d.%m.%Y %H:%M')}</td>
                </tr>
            </table>
        </div>
        
        <div class="report-section">
            <h2 class="section-title">Сравниваемые документы</h2>
            <div class="document-comparison">
                <div class="document-section">
                    <div class="document-title">Базовый документ</div>
                    <div class="content-text">
                        <strong>Название:</strong> {escape(comparison.base_document.title)}<br>
                        <strong>Версия:</strong> {comparison.base_document.version}<br>
                        <strong>Дата загрузки:</strong> {comparison.base_document.upload_date.strftime('%d.%m.%Y %H:%M')}<br>
                        <strong>Статус:</strong> {comparison.base_document.get_status_display()}
                    </div>
                </div>
                <div class="document-section">
                    <div class="document-title">Сравниваемый документ</div>
                    <div class="content-text">
                        <strong>Название:</strong> {escape(comparison.compared_document.title)}<br>
                        <strong>Версия:</strong> {comparison.compared_document.version}<br>
                        <strong>Дата загрузки:</strong> {comparison.compared_document.upload_date.strftime('%d.%m.%Y %H:%M')}<br>
                        <strong>Статус:</strong> {comparison.compared_document.get_status_display()}
                    </div>
                </div>
            </div>
        </div>
        """
        
        return html
    
    def _create_fallback_html(self, report) -> str:
        """Создает базовый HTML для неизвестного формата"""
        return f"""
        <!DOCTYPE html>
        <html lang="ru">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>{escape(report.title)}</title>
            {self.base_styles}
        </head>
        <body>
            <div class="report-container">
                <div class="report-header">
                    <h1 class="report-title">{escape(report.title)}</h1>
                </div>
                
                <div class="highlight-box">
                    <strong>Внимание:</strong> Неподдерживаемый формат отчета: {report.format}
                </div>
                
                <div class="report-section">
                    <h2 class="section-title">Информация об отчете</h2>
                    <p>Название: {escape(report.title)}</p>
                    <p>Формат: {report.format}</p>
                    <p>Дата создания: {report.generated_date.strftime('%d.%m.%Y %H:%M')}</p>
                </div>
            </div>
        </body>
        </html>
        """
    
    def _create_error_html(self, report, error_message: str) -> str:
        """Создает HTML с сообщением об ошибке"""
        return f"""
        <!DOCTYPE html>
        <html lang="ru">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Ошибка - {escape(report.title)}</title>
            {self.base_styles}
        </head>
        <body>
            <div class="report-container">
                <div class="report-header">
                    <h1 class="report-title">Ошибка просмотра отчета</h1>
                </div>
                
                <div class="highlight-box" style="background: #f8d7da; border-color: #f5c6cb; color: #721c24;">
                    <strong>Ошибка:</strong> {escape(error_message)}
                </div>
                
                <div class="report-section">
                    <h2 class="section-title">Информация об отчете</h2>
                    <p>Название: {escape(report.title)}</p>
                    <p>Формат: {report.format}</p>
                    <p>Дата создания: {report.generated_date.strftime('%d.%m.%Y %H:%M')}</p>
                </div>
            </div>
        </body>
        </html>
        """
