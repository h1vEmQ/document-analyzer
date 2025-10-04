"""
Сервисы для генерации отчетов и отправки email
"""
import os
import logging
from typing import Dict, List, Any, Optional
from datetime import datetime
from io import BytesIO
from django.conf import settings
from django.core.mail import EmailMultiAlternatives
from django.template.loader import render_to_string
from django.utils import timezone
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.platypus.flowables import HRFlowable
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from analysis.models import Comparison, Change
from .models import Report, ReportTemplate, EmailNotification
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import json

logger = logging.getLogger(__name__)


class PDFReportGeneratorService:
    """
    Сервис для генерации PDF отчетов
    """
    
    def __init__(self):
        self.page_size = A4
        self.margin = inch
        self.styles = getSampleStyleSheet()
        self._setup_fonts()
        self._setup_custom_styles()
    
    def _setup_fonts(self):
        """Настройка шрифтов для поддержки кириллицы"""
        try:
            # Используем простой подход - встроенные шрифты ReportLab
            # Проблема в том, что ReportLab по умолчанию не поддерживает кириллицу
            # Но мы можем использовать UnicodeCIDFont для поддержки Unicode
            
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont
            
            # Регистрируем шрифт с поддержкой Unicode
            pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
            self.font_name = 'STSong-Light'
            
            logger.info("Зарегистрирован шрифт с поддержкой Unicode: STSong-Light")
            
        except Exception as e:
            logger.error(f"Ошибка при настройке шрифтов: {e}")
            # Fallback на стандартный шрифт
            self.font_name = 'Helvetica'
    
    def _setup_custom_styles(self):
        """Настройка пользовательских стилей для PDF"""
        
        # Стиль для заголовка отчета
        if 'ReportTitle' not in self.styles:
            self.styles.add(ParagraphStyle(
                name='ReportTitle',
                parent=self.styles['Heading1'],
                fontName=self.font_name,
                fontSize=18,
                spaceAfter=30,
                alignment=1,  # Центрирование
                textColor=colors.darkblue
            ))
        
        # Стиль для подзаголовков
        if 'SectionHeader' not in self.styles:
            self.styles.add(ParagraphStyle(
                name='SectionHeader',
                parent=self.styles['Heading2'],
                fontName=self.font_name,
                fontSize=14,
                spaceAfter=12,
                spaceBefore=20,
                textColor=colors.darkblue
            ))
        
        # Стиль для обычного текста
        if 'CustomBodyText' not in self.styles:
            self.styles.add(ParagraphStyle(
                name='CustomBodyText',
                parent=self.styles['Normal'],
                fontName=self.font_name,
                fontSize=10,
                spaceAfter=6,
                leftIndent=0
            ))
        
        # Стиль для изменений
        if 'ChangeText' not in self.styles:
            self.styles.add(ParagraphStyle(
                name='ChangeText',
                parent=self.styles['Normal'],
                fontName=self.font_name,
                fontSize=9,
                spaceAfter=4,
                leftIndent=20,
                rightIndent=20
            ))
    
    def generate_comparison_report(self, comparison: Comparison) -> bytes:
        """
        Генерирует PDF отчет на основе сравнения документов
        """
        try:
            buffer = BytesIO()
            doc = SimpleDocTemplate(
                buffer,
                pagesize=self.page_size,
                rightMargin=self.margin,
                leftMargin=self.margin,
                topMargin=self.margin,
                bottomMargin=self.margin
            )
            
            # Собираем элементы отчета
            story = []
            
            # Заголовок отчета
            story.extend(self._create_report_header(comparison))
            story.append(Spacer(1, 20))
            
            # Сводка изменений
            story.extend(self._create_summary_section(comparison))
            story.append(Spacer(1, 20))
            
            # Детальные изменения
            story.extend(self._create_changes_section(comparison))
            
            # Метаданные
            story.append(PageBreak())
            story.extend(self._create_metadata_section(comparison))
            
            # Генерируем PDF
            doc.build(story)
            
            pdf_data = buffer.getvalue()
            buffer.close()
            
            logger.info(f"PDF отчет для сравнения {comparison.id} сгенерирован успешно")
            return pdf_data
            
        except Exception as e:
            logger.error(f"Ошибка при генерации PDF отчета для сравнения {comparison.id}: {str(e)}")
            raise
    
    def _create_report_header(self, comparison: Comparison) -> List:
        """Создает заголовок отчета"""
        elements = []
        
        # Название отчета
        title = f"Document Changes Report"
        elements.append(Paragraph(title, self.styles['ReportTitle']))
        
        # Информация о документах
        base_doc = comparison.base_document
        compared_doc = comparison.compared_document
        
        info_text = f"""
        <b>Base Document:</b> {base_doc.title} (v{base_doc.version})<br/>
        <b>Compared Document:</b> {compared_doc.title} (v{compared_doc.version})<br/>
        <b>Comparison Date:</b> {comparison.created_date.strftime('%d.%m.%Y %H:%M')}<br/>
        <b>User:</b> {comparison.user.username}
        """
        
        elements.append(Paragraph(info_text, self.styles['CustomBodyText']))
        elements.append(HRFlowable(width="100%", thickness=1, lineCap='round', color=colors.grey))
        
        return elements
    
    def _create_summary_section(self, comparison: Comparison) -> List:
        """Создает раздел сводки изменений"""
        elements = []
        
        elements.append(Paragraph("Changes Summary", self.styles['SectionHeader']))
        
        # Статистика изменений
        summary = comparison.changes_summary or {}
        total_changes = summary.get('total', 0)
        added = summary.get('added', 0)
        removed = summary.get('removed', 0)
        modified = summary.get('modified', 0)
        
        summary_data = [
            ['Change Type', 'Count'],
            ['Total Changes', str(total_changes)],
            ['Added', str(added)],
            ['Removed', str(removed)],
            ['Modified', str(modified)]
        ]
        
        summary_table = Table(summary_data, colWidths=[3*inch, 1.5*inch])
        # Определяем правильные названия шрифтов
        if self.font_name == 'STSong-Light':
            header_font = 'STSong-Light'
            body_font = 'STSong-Light'
        else:
            header_font = 'Helvetica-Bold'
            body_font = 'Helvetica'
        
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), header_font),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('FONTNAME', (0, 1), (-1, -1), body_font),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        elements.append(summary_table)
        
        # Время обработки
        if comparison.processing_time:
            processing_text = f"<b>Analysis Time:</b> {comparison.processing_time:.2f} seconds"
            elements.append(Paragraph(processing_text, self.styles['CustomBodyText']))
        
        return elements
    
    def _create_changes_section(self, comparison: Comparison) -> List:
        """Создает раздел детальных изменений"""
        elements = []
        
        elements.append(Paragraph("Detailed Changes", self.styles['SectionHeader']))
        
        # Группируем изменения по типам
        changes_by_type = {}
        for change in comparison.changes.all():
            change_type = change.change_type
            if change_type not in changes_by_type:
                changes_by_type[change_type] = []
            changes_by_type[change_type].append(change)
        
        # Создаем разделы для каждого типа изменений
        type_labels = {
            'added': 'Added Elements',
            'removed': 'Removed Elements',
            'modified': 'Modified Elements'
        }
        
        for change_type, changes in changes_by_type.items():
            if changes:
                label = type_labels.get(change_type, f"Changes of type '{change_type}'")
                elements.append(Paragraph(label, self.styles['SectionHeader']))
                
                for i, change in enumerate(changes, 1):
                    change_text = f"""
                    <b>{i}.</b> <b>Section:</b> {change.section}<br/>
                    <b>Type:</b> {change.change_type}<br/>
                    <b>Location:</b> {change.location}<br/>
                    <b>Description:</b> {change.new_value[:200]}{'...' if len(change.new_value) > 200 else ''}
                    """
                    
                    elements.append(Paragraph(change_text, self.styles['ChangeText']))
                    elements.append(Spacer(1, 8))
        
        return elements
    
    def _create_metadata_section(self, comparison: Comparison) -> List:
        """Создает раздел метаданных"""
        elements = []
        
        elements.append(Paragraph("Document Metadata", self.styles['SectionHeader']))
        
        # Метаданные базового документа
        base_doc = comparison.base_document
        base_metadata = base_doc.metadata or {}
        
        base_info = f"""
        <b>Base Document:</b><br/>
        • Title: {base_doc.title}<br/>
        • File: {base_doc.filename}<br/>
        • Size: {base_doc.get_file_size_mb()} MB<br/>
        • Upload Date: {base_doc.upload_date.strftime('%d.%m.%Y %H:%M')}<br/>
        • Author: {base_metadata.get('author', 'Not specified')}<br/>
        • Created: {base_metadata.get('created', 'Not specified')}<br/>
        • Modified: {base_metadata.get('modified', 'Not specified')}
        """
        
        elements.append(Paragraph(base_info, self.styles['CustomBodyText']))
        elements.append(Spacer(1, 15))
        
        # Метаданные сравниваемого документа
        compared_doc = comparison.compared_document
        compared_metadata = compared_doc.metadata or {}
        
        compared_info = f"""
        <b>Compared Document:</b><br/>
        • Title: {compared_doc.title}<br/>
        • File: {compared_doc.filename}<br/>
        • Size: {compared_doc.get_file_size_mb()} MB<br/>
        • Upload Date: {compared_doc.upload_date.strftime('%d.%m.%Y %H:%M')}<br/>
        • Author: {compared_metadata.get('author', 'Not specified')}<br/>
        • Created: {compared_metadata.get('created', 'Not specified')}<br/>
        • Modified: {compared_metadata.get('modified', 'Not specified')}
        """
        
        elements.append(Paragraph(compared_info, self.styles['CustomBodyText']))
        elements.append(Spacer(1, 15))
        
        # Информация о сравнении
        comparison_info = f"""
        <b>Comparison Information:</b><br/>
        • Created Date: {comparison.created_date.strftime('%d.%m.%Y %H:%M')}<br/>
        • Completed Date: {comparison.completed_date.strftime('%d.%m.%Y %H:%M') if comparison.completed_date else 'Not completed'}<br/>
        • Status: {comparison.get_status_display()}<br/>
        • User: {comparison.user.username}<br/>
        • Analysis Time: {comparison.processing_time:.2f} seconds
        """
        
        elements.append(Paragraph(comparison_info, self.styles['CustomBodyText']))
        
        return elements


class EmailReportService:
    """
    Сервис для отправки отчетов по email
    """
    
    def __init__(self):
        self.from_email = getattr(settings, 'DEFAULT_FROM_EMAIL', 'noreply@wara.local')
        self.smtp_enabled = getattr(settings, 'EMAIL_USE_TLS', False)
    
    def send_report_email(self, report: Report, recipient_email: str, custom_message: str = "") -> Dict[str, Any]:
        """
        Отправляет отчет по email
        """
        try:
            # Создаем уведомление
            notification = EmailNotification.objects.create(
                report=report,
                recipient_email=recipient_email,
                subject=f"Отчет об изменениях: {report.title}",
                message=custom_message or "Отчет об изменениях документов",
                status='pending'
            )
            
            # Генерируем содержимое email
            subject = f"Отчет об изменениях: {report.title}"
            
            # HTML версия письма
            html_content = self._generate_email_html(report, custom_message)
            
            # Текстовая версия письма
            text_content = self._generate_email_text(report, custom_message)
            
            # Создаем email
            email = EmailMultiAlternatives(
                subject=subject,
                body=text_content,
                from_email=self.from_email,
                to=[recipient_email]
            )
            
            email.attach_alternative(html_content, "text/html")
            
            # Прикрепляем PDF файл
            if report.file and report.file.name:
                email.attach_file(report.file.path)
            
            # Отправляем email
            if self.smtp_enabled:
                email.send()
                notification.status = 'sent'
                notification.sent_date = timezone.now()
                notification.save()
                
                logger.info(f"Email отчет {report.id} отправлен на {recipient_email}")
                
                return {
                    'success': True,
                    'message': 'Отчет успешно отправлен по email',
                    'notification_id': notification.id
                }
            else:
                # В режиме разработки просто логируем
                logger.info(f"Email не отправлен (SMTP отключен): {subject}")
                notification.status = 'failed'
                notification.save()
                
                return {
                    'success': False,
                    'message': 'SMTP не настроен. Email не отправлен.',
                    'notification_id': notification.id
                }
                
        except Exception as e:
            logger.error(f"Ошибка при отправке email отчета {report.id}: {str(e)}")
            
            if 'notification' in locals():
                notification.status = 'failed'
                notification.save()
            
            return {
                'success': False,
                'message': f'Ошибка при отправке email: {str(e)}',
                'notification_id': notification.id if 'notification' in locals() else None
            }
    
    def _generate_email_html(self, report: Report, custom_message: str) -> str:
        """Генерирует HTML содержимое email"""
        
        comparison = report.comparison
        summary = comparison.changes_summary or {}
        
        html_content = f"""
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                .header {{ background-color: #f8f9fa; padding: 20px; border-radius: 5px; }}
                .summary {{ background-color: #e9ecef; padding: 15px; border-radius: 5px; margin: 20px 0; }}
                .footer {{ margin-top: 30px; padding-top: 20px; border-top: 1px solid #dee2e6; }}
                .highlight {{ color: #007bff; font-weight: bold; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h2>Отчет об изменениях документов</h2>
                <p><strong>Отчет:</strong> {report.title}</p>
                <p><strong>Дата генерации:</strong> {report.generated_date.strftime('%d.%m.%Y %H:%M')}</p>
            </div>
            
            <div class="summary">
                <h3>Сводка изменений</h3>
                <p><span class="highlight">Всего изменений:</span> {summary.get('total', 0)}</p>
                <p><span class="highlight">Добавлено:</span> {summary.get('added', 0)}</p>
                <p><span class="highlight">Удалено:</span> {summary.get('removed', 0)}</p>
                <p><span class="highlight">Изменено:</span> {summary.get('modified', 0)}</p>
            </div>
            
            <div>
                <h3>Документы</h3>
                <p><strong>Базовый документ:</strong> {comparison.base_document.title}</p>
                <p><strong>Сравниваемый документ:</strong> {comparison.compared_document.title}</p>
            </div>
            
            {f'<div><h3>Сообщение</h3><p>{custom_message}</p></div>' if custom_message else ''}
            
            <div class="footer">
                <p>Подробный отчет прикреплен к письму в формате PDF.</p>
                <p>Сгенерировано системой Document analyzer ({timezone.now().strftime('%d.%m.%Y %H:%M')})</p>
            </div>
        </body>
        </html>
        """
        
        return html_content
    
    def _generate_email_text(self, report: Report, custom_message: str) -> str:
        """Генерирует текстовое содержимое email"""
        
        comparison = report.comparison
        summary = comparison.changes_summary or {}
        
        text_content = f"""
Отчет об изменениях документов

Отчет: {report.title}
Дата генерации: {report.generated_date.strftime('%d.%m.%Y %H:%M')}

СВОДКА ИЗМЕНЕНИЙ:
- Всего изменений: {summary.get('total', 0)}
- Добавлено: {summary.get('added', 0)}
- Удалено: {summary.get('removed', 0)}
- Изменено: {summary.get('modified', 0)}

ДОКУМЕНТЫ:
- Базовый документ: {comparison.base_document.title}
- Сравниваемый документ: {comparison.compared_document.title}

{f'СООБЩЕНИЕ:\n{custom_message}\n' if custom_message else ''}

Подробный отчет прикреплен к письму в формате PDF.

Сгенерировано системой Document analyzer ({timezone.now().strftime('%d.%m.%Y %H:%M')})
        """
        
        return text_content


class ReportTemplateService:
    """
    Сервис для работы с шаблонами отчетов
    """
    
    def get_default_template(self) -> ReportTemplate:
        """Получает шаблон по умолчанию"""
        template, created = ReportTemplate.objects.get_or_create(
            is_default=True,
            defaults={
                'name': 'Шаблон по умолчанию',
                'template_content': self._get_default_template_content(),
                'is_default': True
            }
        )
        return template
    
    def _get_default_template_content(self) -> str:
        """Возвращает содержимое шаблона по умолчанию"""
        return """
        <h1>Отчет об изменениях документов</h1>
        
        <h2>Информация о сравнении</h2>
        <p><strong>Базовый документ:</strong> {{ comparison.base_document.title }}</p>
        <p><strong>Сравниваемый документ:</strong> {{ comparison.compared_document.title }}</p>
        <p><strong>Дата сравнения:</strong> {{ comparison.created_date }}</p>
        
        <h2>Сводка изменений</h2>
        <p><strong>Всего изменений:</strong> {{ summary.total }}</p>
        <p><strong>Добавлено:</strong> {{ summary.added }}</p>
        <p><strong>Удалено:</strong> {{ summary.removed }}</p>
        <p><strong>Изменено:</strong> {{ summary.modified }}</p>
        
        <h2>Детальные изменения</h2>
        {% for change in changes %}
            <div class="change-item">
                <h3>{{ change.section }}</h3>
                <p><strong>Тип:</strong> {{ change.change_type }}</p>
                <p><strong>Описание:</strong> {{ change.new_value }}</p>
            </div>
        {% endfor %}
        """


class DOCXReportGeneratorService:
    """
    Сервис для генерации DOCX отчетов
    """
    
    def __init__(self):
        self.doc = None
    
    def generate_comparison_report(self, comparison: Comparison) -> bytes:
        """
        Генерирует DOCX отчет для сравнения документов
        
        Args:
            comparison: Объект сравнения
            
        Returns:
            bytes: Содержимое DOCX файла
        """
        return self.generate_report(comparison)
    
    def generate_report(self, comparison: Comparison, options: Dict[str, Any] = None) -> bytes:
        """
        Генерирует DOCX отчет для сравнения документов
        
        Args:
            comparison: Объект сравнения
            options: Дополнительные опции генерации
            
        Returns:
            bytes: Содержимое DOCX файла
        """
        try:
            # Создаем новый документ
            self.doc = DocxDocument()
            
            # Настраиваем стили
            self._setup_styles()
            
            # Создаем отчет
            self._create_report_header(comparison)
            self._create_summary_section(comparison)
            self._create_changes_section(comparison)
            self._create_metadata_section(comparison)
            
            # Сохраняем в байты
            buffer = BytesIO()
            self.doc.save(buffer)
            buffer.seek(0)
            
            return buffer.getvalue()
            
        except Exception as e:
            logger.error(f"Error generating DOCX report for comparison {comparison.id}: {e}")
            raise
    
    def _setup_styles(self):
        """Настройка стилей документа"""
        # Основные стили уже настроены в python-docx по умолчанию
        pass
    
    def _create_report_header(self, comparison: Comparison):
        """Создание заголовка отчета"""
        # Заголовок
        title = self.doc.add_heading(f'Document Comparison Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Подзаголовок
        subtitle = self.doc.add_heading(f'Comparison: {comparison.title}', level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о документах
        doc_info = self.doc.add_paragraph()
        doc_info.add_run(f'Base Document: ').bold = True
        doc_info.add_run(f'{comparison.base_document.title} (v{comparison.base_document.version})')
        doc_info.add_run('\n')
        doc_info.add_run(f'Compared Document: ').bold = True
        doc_info.add_run(f'{comparison.compared_document.title} (v{comparison.compared_document.version})')
        
        # Дата анализа
        date_para = self.doc.add_paragraph()
        date_para.add_run(f'Analysis Date: ').bold = True
        date_para.add_run(f'{comparison.completed_date.strftime("%Y-%m-%d %H:%M") if comparison.completed_date else "N/A"}')
        
        # Разделитель
        self.doc.add_paragraph('─' * 80)
    
    def _create_summary_section(self, comparison: Comparison):
        """Создание раздела сводки"""
        self.doc.add_heading('Summary', level=1)
        
        if comparison.changes_summary:
            summary = comparison.changes_summary
            
            # Создаем таблицу сводки
            table = self.doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Заголовки таблицы
            table.cell(0, 0).text = 'Metric'
            table.cell(0, 1).text = 'Count'
            
            # Заполняем таблицу
            table.cell(1, 0).text = 'Total Changes'
            table.cell(1, 1).text = str(summary.get('total', 0))
            
            table.cell(2, 0).text = 'Added'
            table.cell(2, 1).text = str(summary.get('added', 0))
            
            table.cell(3, 0).text = 'Modified'
            table.cell(3, 1).text = str(summary.get('modified', 0))
            
            # Выделяем заголовки
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
        else:
            self.doc.add_paragraph('No summary data available.')
        
        self.doc.add_paragraph()
    
    def _create_changes_section(self, comparison: Comparison):
        """Создание раздела детальных изменений"""
        self.doc.add_heading('Detailed Changes', level=1)
        
        changes = comparison.changes.all().order_by('id')
        
        if changes:
            for i, change in enumerate(changes, 1):
                # Заголовок изменения
                change_heading = self.doc.add_heading(f'Change {i}: {change.get_change_type_display()}', level=2)
                
                # Информация об изменении
                info_para = self.doc.add_paragraph()
                info_para.add_run(f'Location: ').bold = True
                info_para.add_run(f'{change.get_location_display()}')
                info_para.add_run(' | ')
                info_para.add_run(f'Section: ').bold = True
                info_para.add_run(f'{change.section or "N/A"}')
                
                # Старое значение
                if change.old_value:
                    old_para = self.doc.add_paragraph()
                    old_para.add_run('Previous Value: ').bold = True
                    old_para.add_run(change.old_value[:500] + ('...' if len(change.old_value) > 500 else ''))
                
                # Новое значение
                if change.new_value:
                    new_para = self.doc.add_paragraph()
                    new_para.add_run('New Value: ').bold = True
                    new_para.add_run(change.new_value[:500] + ('...' if len(change.new_value) > 500 else ''))
                
                # Уверенность
                confidence_para = self.doc.add_paragraph()
                confidence_para.add_run(f'Confidence: ').bold = True
                confidence_para.add_run(f'{change.confidence:.2f}')
                
                self.doc.add_paragraph('─' * 60)
        else:
            self.doc.add_paragraph('No changes found.')
    
    def _create_metadata_section(self, comparison: Comparison):
        """Создание раздела метаданных"""
        self.doc.add_heading('Analysis Metadata', level=1)
        
        # Создаем таблицу метаданных
        table = self.doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        
        # Заполняем таблицу
        metadata = [
            ('Comparison ID', str(comparison.id)),
            ('Base Document ID', str(comparison.base_document.id)),
            ('Compared Document ID', str(comparison.compared_document.id)),
            ('Status', comparison.get_status_display()),
            ('Created Date', comparison.created_date.strftime('%Y-%m-%d %H:%M')),
            ('Completed Date', comparison.completed_date.strftime('%Y-%m-%d %H:%M') if comparison.completed_date else 'N/A'),
        ]
        
        for i, (key, value) in enumerate(metadata):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = value
            
            # Выделяем заголовки
            for paragraph in table.cell(i, 0).paragraphs:
                for run in paragraph.runs:
                    run.bold = True


class AutoReportGeneratorService:
    """
    Сервис для автоматической генерации отчетов после анализа
    """
    
    def __init__(self):
        self.pdf_generator = PDFReportGeneratorService()
        self.docx_generator = DOCXReportGeneratorService()
    
    def generate_auto_reports(self, comparison: Comparison) -> Dict[str, Any]:
        """
        Автоматически генерирует отчет в формате, установленном в настройках приложения
        
        Args:
            comparison: Объект сравнения
            
        Returns:
            Dict с результатами генерации отчета
        """
        # Получаем формат отчета из настроек приложения
        from settings.models import ApplicationSettings
        settings = ApplicationSettings.get_settings()
        report_format = settings.default_report_format
        
        results = {
            'pdf_report': None,
            'docx_report': None,
            'errors': []
        }
        
        try:
            # Генерируем отчет в выбранном формате
            if report_format == 'docx':
                try:
                    docx_content = self.docx_generator.generate_comparison_report(comparison)
                    docx_report = self._save_report(comparison, docx_content, 'docx', f'Auto-generated DOCX Report')
                    results['docx_report'] = docx_report
                    logger.info(f"Auto-generated DOCX report for comparison {comparison.id}")
                except Exception as e:
                    error_msg = f"Failed to generate DOCX report: {str(e)}"
                    results['errors'].append(error_msg)
                    logger.error(error_msg)
            else:  # pdf
                try:
                    pdf_content = self.pdf_generator.generate_comparison_report(comparison)
                    pdf_report = self._save_report(comparison, pdf_content, 'pdf', f'Auto-generated PDF Report')
                    results['pdf_report'] = pdf_report
                    logger.info(f"Auto-generated PDF report for comparison {comparison.id}")
                except Exception as e:
                    error_msg = f"Failed to generate PDF report: {str(e)}"
                    results['errors'].append(error_msg)
                    logger.error(error_msg)
            
            return results
            
        except Exception as e:
            error_msg = f"Failed to generate auto reports for comparison {comparison.id}: {str(e)}"
            results['errors'].append(error_msg)
            logger.error(error_msg)
            return results
    
    def _save_report(self, comparison: Comparison, content: bytes, format_type: str, title: str) -> Report:
        """
        Сохраняет отчет в базу данных с поддержкой версионирования
        
        Args:
            comparison: Объект сравнения
            content: Содержимое отчета
            format_type: Тип формата ('pdf' или 'docx')
            title: Название отчета
            
        Returns:
            Report: Созданный объект отчета
        """
        # Создаем имя файла
        timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
        filename = f"comparison_{comparison.id}_{timestamp}.{format_type}"
        
        # Сохраняем файл
        file_path = os.path.join(settings.MEDIA_ROOT, 'reports', filename)
        os.makedirs(os.path.dirname(file_path), exist_ok=True)
        
        with open(file_path, 'wb') as f:
            f.write(content)
        
        # Проверяем, есть ли уже отчеты для этого сравнения в том же формате
        existing_reports = Report.objects.filter(
            comparison=comparison,
            format=format_type,
            user=comparison.user
        ).order_by('-generated_date')
        
        if existing_reports.exists():
            # Если есть существующие отчеты, создаем новую версию
            latest_report = existing_reports.first()
            version_notes = f"Автоматически созданная версия после анализа сравнения {comparison.id}"
            
            # Создаем новую версию отчета
            report = latest_report.create_new_version(
                new_file=os.path.join('reports', filename),
                version_notes=version_notes
            )
            
            logger.info(f"Created new version {report.version} of report {latest_report.id}")
        else:
            # Если это первый отчет для данного сравнения, создаем корневой отчет
            report = Report.objects.create(
                title=title,
                comparison=comparison,
                user=comparison.user,
                format=format_type,
                file=os.path.join('reports', filename),
                status='ready',
                include_summary=True,
                include_details=True,
                include_tables=True,
                version='1.0',
                is_latest_version=True,
                version_notes='Первая версия отчета',
                generated_date=timezone.now()
            )
            
            logger.info(f"Created new root report {report.id} for comparison {comparison.id}")
        
        return report


class OllamaReportGeneratorService:
    """
    Сервис для генерации отчетов по результатам анализа нейросетью
    """
    
    def __init__(self):
        self.pdf_generator = PDFReportGeneratorService()
        self.docx_generator = DOCXReportGeneratorService()
    
    def generate_ollama_report(self, comparison: Comparison, format: str = None) -> bytes:
        """
        Генерирует отчет по результатам анализа нейросетью
        
        Args:
            comparison: Объект сравнения с результатами анализа
            format: Формат отчета ('pdf' или 'docx'). Если None, берется из настроек
            
        Returns:
            bytes: Содержимое файла отчета
        """
        # Если формат не указан, получаем из настроек
        if format is None:
            from settings.models import ApplicationSettings
            settings = ApplicationSettings.get_settings()
            format = settings.default_report_format
        
        if format.lower() == 'docx':
            return self._generate_docx_ollama_report(comparison)
        else:
            return self._generate_pdf_ollama_report(comparison)
    
    def _generate_pdf_ollama_report(self, comparison: Comparison) -> bytes:
        """Генерирует PDF отчет по результатам анализа нейросетью"""
        buffer = BytesIO()
        
        # Создаем PDF документ
        doc = SimpleDocTemplate(
            buffer,
            pagesize=self.pdf_generator.page_size,
            rightMargin=self.pdf_generator.margin,
            leftMargin=self.pdf_generator.margin,
            topMargin=self.pdf_generator.margin,
            bottomMargin=self.pdf_generator.margin
        )
        
        # Собираем элементы отчета
        elements = []
        
        # Заголовок
        title = Paragraph(f"Отчет анализа нейросетью: {comparison.title}", 
                         self.pdf_generator.styles['Title'])
        elements.append(title)
        elements.append(Spacer(1, 12))
        
        # Информация об анализе
        info_data = [
            ['Базовый документ:', f"{comparison.base_document.title} (v{comparison.base_document.version})"],
            ['Сравниваемый документ:', f"{comparison.compared_document.title} (v{comparison.compared_document.version})"],
            ['Модель нейросети:', comparison.analysis_method or 'Не указана'],
            ['Дата анализа:', comparison.created_date.strftime('%d.%m.%Y %H:%M')],
            ['Статус:', comparison.get_status_display()]
        ]
        
        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSize', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('BACKGROUND', (1, 0), (1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        elements.append(info_table)
        elements.append(Spacer(1, 20))
        
        # Результаты анализа
        if comparison.analysis_result:
            elements.append(Paragraph("Результаты анализа", self.pdf_generator.styles['Heading2']))
            elements.append(Spacer(1, 12))
            
            analysis_data = comparison.analysis_result
            
            # Резюме
            if analysis_data.get('summary'):
                elements.append(Paragraph("Резюме", self.pdf_generator.styles['Heading3']))
                elements.append(Paragraph(analysis_data['summary'], self.pdf_generator.styles['Normal']))
                elements.append(Spacer(1, 12))
            
            # Сходства
            if analysis_data.get('similarities'):
                elements.append(Paragraph("Сходства", self.pdf_generator.styles['Heading3']))
                for similarity in analysis_data['similarities']:
                    elements.append(Paragraph(f"• {similarity}", self.pdf_generator.styles['Normal']))
                elements.append(Spacer(1, 12))
            
            # Различия
            if analysis_data.get('differences'):
                elements.append(Paragraph("Различия", self.pdf_generator.styles['Heading3']))
                for i, difference in enumerate(analysis_data['differences'], 1):
                    elements.append(Paragraph(f"{i}. {difference.get('description', 'Описание отсутствует')}", 
                                            self.pdf_generator.styles['Normal']))
                    
                    if difference.get('location'):
                        elements.append(Paragraph(f"   Место: {difference['location']}", 
                                                self.pdf_generator.styles['Normal']))
                    
                    if difference.get('significance'):
                        significance_colors = {
                            'high': colors.red,
                            'medium': colors.orange,
                            'low': colors.green
                        }
                        color = significance_colors.get(difference['significance'], colors.black)
                        elements.append(Paragraph(f"   Важность: {difference['significance']}", 
                                                self.pdf_generator.styles['Normal']))
                    
                    elements.append(Spacer(1, 6))
            
            # Рекомендации
            if analysis_data.get('recommendations'):
                elements.append(Paragraph("Рекомендации", self.pdf_generator.styles['Heading3']))
                for recommendation in analysis_data['recommendations']:
                    elements.append(Paragraph(f"• {recommendation}", self.pdf_generator.styles['Normal']))
                elements.append(Spacer(1, 12))
            
            # Общая оценка
            if analysis_data.get('overall_assessment'):
                elements.append(Paragraph("Общая оценка", self.pdf_generator.styles['Heading3']))
                elements.append(Paragraph(analysis_data['overall_assessment'], self.pdf_generator.styles['Normal']))
                elements.append(Spacer(1, 12))
            
            # Анализ тональности (если есть)
            if analysis_data.get('base_document_sentiment') and analysis_data.get('compared_document_sentiment'):
                elements.append(PageBreak())
                elements.append(Paragraph("Анализ тональности", self.pdf_generator.styles['Heading2']))
                elements.append(Spacer(1, 12))
                
                # Тональность базового документа
                base_sentiment = analysis_data['base_document_sentiment']
                elements.append(Paragraph(f"Тональность базового документа: {comparison.base_document.title} (v{comparison.base_document.version})", 
                                        self.pdf_generator.styles['Heading3']))
                
                sentiment_data = [
                    ['Параметр', 'Значение'],
                    ['Тональность', base_sentiment.get('sentiment', 'Не определена')],
                    ['Уверенность', str(base_sentiment.get('confidence', 'Не указана'))],
                    ['Резюме', base_sentiment.get('summary', 'Отсутствует')]
                ]
                
                if base_sentiment.get('emotions'):
                    sentiment_data.append(['Эмоции', ', '.join(base_sentiment['emotions'])])
                
                sentiment_table = Table(sentiment_data, colWidths=[2*inch, 4*inch])
                sentiment_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                    ('FONTSize', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                    ('BACKGROUND', (1, 1), (1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                elements.append(sentiment_table)
                elements.append(Spacer(1, 20))
                
                # Тональность сравниваемого документа
                compared_sentiment = analysis_data['compared_document_sentiment']
                elements.append(Paragraph(f"Тональность сравниваемого документа: {comparison.compared_document.title} (v{comparison.compared_document.version})", 
                                        self.pdf_generator.styles['Heading3']))
                
                sentiment_data2 = [
                    ['Параметр', 'Значение'],
                    ['Тональность', compared_sentiment.get('sentiment', 'Не определена')],
                    ['Уверенность', str(compared_sentiment.get('confidence', 'Не указана'))],
                    ['Резюме', compared_sentiment.get('summary', 'Отсутствует')]
                ]
                
                if compared_sentiment.get('emotions'):
                    sentiment_data2.append(['Эмоции', ', '.join(compared_sentiment['emotions'])])
                
                sentiment_table2 = Table(sentiment_data2, colWidths=[2*inch, 4*inch])
                sentiment_table2.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                    ('FONTSize', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                    ('BACKGROUND', (1, 1), (1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                elements.append(sentiment_table2)
                elements.append(Spacer(1, 20))
            
            # Ключевые моменты (если есть)
            if analysis_data.get('base_document_key_points') and analysis_data.get('compared_document_key_points'):
                elements.append(PageBreak())
                elements.append(Paragraph("Ключевые моменты", self.pdf_generator.styles['Heading2']))
                elements.append(Spacer(1, 12))
                
                # Ключевые моменты базового документа
                base_key_points = analysis_data['base_document_key_points']
                elements.append(Paragraph(f"Ключевые моменты: {comparison.base_document.title} (v{comparison.base_document.version})", 
                                        self.pdf_generator.styles['Heading3']))
                
                if base_key_points.get('key_points'):
                    for point in base_key_points['key_points']:
                        importance = point.get('importance', 'medium')
                        importance_colors = {
                            'high': colors.red,
                            'medium': colors.orange,
                            'low': colors.green
                        }
                        color = importance_colors.get(importance, colors.black)
                        
                        elements.append(Paragraph(f"• {point.get('point', 'Точка не указана')}", 
                                                self.pdf_generator.styles['Normal']))
                        if point.get('category'):
                            elements.append(Paragraph(f"  Категория: {point['category']}", 
                                                    self.pdf_generator.styles['Normal']))
                        elements.append(Spacer(1, 6))
                
                elements.append(Spacer(1, 12))
                
                # Ключевые моменты сравниваемого документа
                compared_key_points = analysis_data['compared_document_key_points']
                elements.append(Paragraph(f"Ключевые моменты: {comparison.compared_document.title} (v{comparison.compared_document.version})", 
                                        self.pdf_generator.styles['Heading3']))
                
                if compared_key_points.get('key_points'):
                    for point in compared_key_points['key_points']:
                        importance = point.get('importance', 'medium')
                        importance_colors = {
                            'high': colors.red,
                            'medium': colors.orange,
                            'low': colors.green
                        }
                        color = importance_colors.get(importance, colors.black)
                        
                        elements.append(Paragraph(f"• {point.get('point', 'Точка не указана')}", 
                                                self.pdf_generator.styles['Normal']))
                        if point.get('category'):
                            elements.append(Paragraph(f"  Категория: {point['category']}", 
                                                    self.pdf_generator.styles['Normal']))
                        elements.append(Spacer(1, 6))
        
        # Сырой ответ модели (если есть)
        if comparison.analysis_result and comparison.analysis_result.get('raw_analysis'):
            elements.append(PageBreak())
            elements.append(Paragraph("Сырой ответ модели", self.pdf_generator.styles['Heading2']))
            elements.append(Spacer(1, 12))
            
            raw_analysis = comparison.analysis_result['raw_analysis']
            # Ограничиваем длину для PDF
            if len(raw_analysis) > 2000:
                raw_analysis = raw_analysis[:2000] + "... (текст обрезан)"
            
            elements.append(Paragraph(raw_analysis, self.pdf_generator.styles['Normal']))
        
        # Собираем PDF
        doc.build(elements)
        
        # Получаем содержимое
        pdf_content = buffer.getvalue()
        buffer.close()
        
        return pdf_content
    
    def _generate_docx_ollama_report(self, comparison: Comparison) -> bytes:
        """Генерирует DOCX отчет по результатам анализа нейросетью"""
        doc = DocxDocument()
        
        # Заголовок
        title = doc.add_heading(f'Отчет анализа нейросетью: {comparison.title}', 0)
        
        # Информация об анализе
        doc.add_heading('Информация об анализе', level=1)
        
        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Базовый документ:', f"{comparison.base_document.title} (v{comparison.base_document.version})"),
            ('Сравниваемый документ:', f"{comparison.compared_document.title} (v{comparison.compared_document.version})"),
            ('Модель нейросети:', comparison.analysis_method or 'Не указана'),
            ('Дата анализа:', comparison.created_date.strftime('%d.%m.%Y %H:%M')),
            ('Статус:', comparison.get_status_display())
        ]
        
        for i, (key, value) in enumerate(info_data):
            info_table.cell(i, 0).text = key
            info_table.cell(i, 1).text = str(value)
        
        # Результаты анализа
        if comparison.analysis_result:
            doc.add_heading('Результаты анализа', level=1)
            
            analysis_data = comparison.analysis_result
            
            # Резюме
            if analysis_data.get('summary'):
                doc.add_heading('Резюме', level=2)
                doc.add_paragraph(analysis_data['summary'])
            
            # Сходства
            if analysis_data.get('similarities'):
                doc.add_heading('Сходства', level=2)
                for similarity in analysis_data['similarities']:
                    doc.add_paragraph(f'• {similarity}', style='List Bullet')
            
            # Различия
            if analysis_data.get('differences'):
                doc.add_heading('Различия', level=2)
                for i, difference in enumerate(analysis_data['differences'], 1):
                    p = doc.add_paragraph(f'{i}. {difference.get("description", "Описание отсутствует")}')
                    
                    if difference.get('location'):
                        doc.add_paragraph(f'   Место: {difference["location"]}', style='List Bullet')
                    
                    if difference.get('significance'):
                        doc.add_paragraph(f'   Важность: {difference["significance"]}', style='List Bullet')
            
            # Рекомендации
            if analysis_data.get('recommendations'):
                doc.add_heading('Рекомендации', level=2)
                for recommendation in analysis_data['recommendations']:
                    doc.add_paragraph(f'• {recommendation}', style='List Bullet')
            
            # Общая оценка
            if analysis_data.get('overall_assessment'):
                doc.add_heading('Общая оценка', level=2)
                doc.add_paragraph(analysis_data['overall_assessment'])
            
            # Анализ тональности
            if analysis_data.get('base_document_sentiment') and analysis_data.get('compared_document_sentiment'):
                doc.add_page_break()
                doc.add_heading('Анализ тональности', level=1)
                
                # Тональность базового документа
                base_sentiment = analysis_data['base_document_sentiment']
                doc.add_heading(f'Тональность: {comparison.base_document.title} (v{comparison.base_document.version})', level=2)
                
                sentiment_table = doc.add_table(rows=4, cols=2)
                sentiment_table.style = 'Table Grid'
                
                sentiment_data = [
                    ('Тональность', base_sentiment.get('sentiment', 'Не определена')),
                    ('Уверенность', str(base_sentiment.get('confidence', 'Не указана'))),
                    ('Резюме', base_sentiment.get('summary', 'Отсутствует'))
                ]
                
                if base_sentiment.get('emotions'):
                    sentiment_data.append(('Эмоции', ', '.join(base_sentiment['emotions'])))
                
                for i, (key, value) in enumerate(sentiment_data):
                    sentiment_table.cell(i, 0).text = key
                    sentiment_table.cell(i, 1).text = str(value)
                
                # Тональность сравниваемого документа
                compared_sentiment = analysis_data['compared_document_sentiment']
                doc.add_heading(f'Тональность: {comparison.compared_document.title} (v{comparison.compared_document.version})', level=2)
                
                sentiment_table2 = doc.add_table(rows=4, cols=2)
                sentiment_table2.style = 'Table Grid'
                
                sentiment_data2 = [
                    ('Тональность', compared_sentiment.get('sentiment', 'Не определена')),
                    ('Уверенность', str(compared_sentiment.get('confidence', 'Не указана'))),
                    ('Резюме', compared_sentiment.get('summary', 'Отсутствует'))
                ]
                
                if compared_sentiment.get('emotions'):
                    sentiment_data2.append(('Эмоции', ', '.join(compared_sentiment['emotions'])))
                
                for i, (key, value) in enumerate(sentiment_data2):
                    sentiment_table2.cell(i, 0).text = key
                    sentiment_table2.cell(i, 1).text = str(value)
            
            # Ключевые моменты
            if analysis_data.get('base_document_key_points') and analysis_data.get('compared_document_key_points'):
                doc.add_page_break()
                doc.add_heading('Ключевые моменты', level=1)
                
                # Ключевые моменты базового документа
                base_key_points = analysis_data['base_document_key_points']
                doc.add_heading(f'Ключевые моменты: {comparison.base_document.title} (v{comparison.base_document.version})', level=2)
                
                if base_key_points.get('key_points'):
                    for point in base_key_points['key_points']:
                        p = doc.add_paragraph(f'• {point.get("point", "Точка не указана")}', style='List Bullet')
                        if point.get('category'):
                            doc.add_paragraph(f'  Категория: {point["category"]}')
                
                # Ключевые моменты сравниваемого документа
                compared_key_points = analysis_data['compared_document_key_points']
                doc.add_heading(f'Ключевые моменты: {comparison.compared_document.title} (v{comparison.compared_document.version})', level=2)
                
                if compared_key_points.get('key_points'):
                    for point in compared_key_points['key_points']:
                        p = doc.add_paragraph(f'• {point.get("point", "Точка не указана")}', style='List Bullet')
                        if point.get('category'):
                            doc.add_paragraph(f'  Категория: {point["category"]}')
        
        # Сырой ответ модели
        if comparison.analysis_result and comparison.analysis_result.get('raw_analysis'):
            doc.add_page_break()
            doc.add_heading('Сырой ответ модели', level=1)
            
            raw_analysis = comparison.analysis_result['raw_analysis']
            # Ограничиваем длину для DOCX
            if len(raw_analysis) > 5000:
                raw_analysis = raw_analysis[:5000] + "... (текст обрезан)"
            
            doc.add_paragraph(raw_analysis)
        
        # Сохраняем в BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def save_ollama_report(self, comparison: Comparison, format: str = None) -> Report:
        """
        Сохраняет отчет по результатам анализа нейросетью в базу данных
        
        Args:
            comparison: Объект сравнения с результатами анализа
            format: Формат отчета ('pdf' или 'docx'). Если None, берется из настроек
            
        Returns:
            Report: Созданный объект отчета
        """
        # Если формат не указан, получаем из настроек
        if format is None:
            from settings.models import ApplicationSettings
            settings = ApplicationSettings.get_settings()
            format = settings.default_report_format
        
        # Генерируем содержимое отчета
        report_content = self.generate_ollama_report(comparison, format)
        
        # Определяем расширение файла
        file_extension = 'pdf' if format.lower() == 'pdf' else 'docx'
        
        # Создаем объект отчета
        report = Report.objects.create(
            user=comparison.user,
            comparison=comparison,
            title=f"Отчет анализа нейросетью: {comparison.title}",
            format=format,
            file=None,  # Будет установлено через ContentFile
            template_used='ollama_ai_analysis',
            status='ready',
            include_tables=False,
            version='1.0',
            is_latest_version=True,
            version_notes='Отчет анализа нейросетью',
            generated_date=timezone.now()
        )
        
        # Сохраняем файл
        from django.core.files.base import ContentFile
        filename = f"ollama_analysis_report_{comparison.id}_{timezone.now().strftime('%Y%m%d_%H%M%S')}.{file_extension}"
        report.file.save(filename, ContentFile(report_content), save=True)
        
        logger.info(f"Created Ollama analysis report {report.id} for comparison {comparison.id}")
        
        return report
