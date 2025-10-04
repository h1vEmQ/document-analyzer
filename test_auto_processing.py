#!/usr/bin/env python
"""
Тестирование автоматической обработки документов при загрузке
"""

import os
import sys
import django

# Настройка Django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'wara_project.settings')
django.setup()

from django.test import Client
from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from documents.models import Document as DjangoDocument
from documents.services import DocumentParserService
from docx import Document as DocxDocument
from io import BytesIO

User = get_user_model()

def create_test_docx():
    """Создание тестового .docx файла"""
    doc = DocxDocument()
    
    # Добавляем заголовок
    doc.add_heading('Тестовый документ для автоматической обработки', 0)
    
    # Добавляем параграф
    doc.add_paragraph('Это тестовый документ, созданный для проверки автоматической обработки.')
    
    # Добавляем раздел
    doc.add_heading('Раздел 1: Основная информация', level=1)
    doc.add_paragraph('Содержимое первого раздела с важной информацией.')
    
    # Добавляем таблицу
    doc.add_heading('Таблица данных', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Заполняем таблицу
    table.cell(0, 0).text = 'Параметр'
    table.cell(0, 1).text = 'Значение'
    table.cell(0, 2).text = 'Единица измерения'
    
    table.cell(1, 0).text = 'Скорость'
    table.cell(1, 1).text = '120'
    table.cell(1, 2).text = 'км/ч'
    
    table.cell(2, 0).text = 'Время'
    table.cell(2, 1).text = '2.5'
    table.cell(2, 2).text = 'часа'
    
    # Добавляем еще один раздел
    doc.add_heading('Раздел 2: Дополнительная информация', level=1)
    doc.add_paragraph('Второй раздел с дополнительными данными для тестирования.')
    
    # Сохраняем в байты
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def test_auto_processing():
    """Тестирование автоматической обработки"""
    print("🧪 Тестирование автоматической обработки документов")
    print("=" * 60)
    
    # Получаем или создаем тестового пользователя
    user, created = User.objects.get_or_create(
        username='testuser',
        defaults={
            'email': 'test@example.com',
            'first_name': 'Тест',
            'last_name': 'Пользователь',
            'role': 'user'
        }
    )
    if created:
        user.set_password('testpass123')
        user.save()
        print(f"    ✅ Создан пользователь: {user.username}")
    else:
        print(f"    ✅ Используется существующий пользователь: {user.username}")
    
    # Создаем тестовый .docx файл
    print("\n📄 Создание тестового документа...")
    docx_content = create_test_docx()
    test_file = SimpleUploadedFile(
        "test_auto_processing.docx",
        docx_content,
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    print(f"    ✅ Создан тестовый файл: {test_file.name} ({len(docx_content)} байт)")
    
    # Создаем клиент для тестирования
    client = Client()
    client.force_login(user)
    
    # Тест 1: Загрузка документа через форму
    print("\n🔍 Тест 1: Загрузка документа через форму")
    response = client.post('/upload/', {
        'title': 'Автоматическая обработка тест',
        'file': test_file
    })
    
    if response.status_code == 302:  # Редирект после успешной загрузки
        print("    ✅ Документ успешно загружен")
        
        # Получаем последний загруженный документ
        latest_doc = DjangoDocument.objects.filter(user=user).order_by('-upload_date').first()
        if latest_doc:
            print(f"    📝 Документ ID: {latest_doc.id}")
            print(f"    📝 Название: {latest_doc.title}")
            print(f"    📝 Статус: {latest_doc.status}")
            
            # Проверяем, был ли документ автоматически обработан
            if latest_doc.status == 'processed':
                print("    ✅ Документ автоматически обработан!")
                
                # Проверяем содержимое
                if latest_doc.content_text:
                    print(f"    📄 Извлеченный текст: {len(latest_doc.content_text)} символов")
                
                # Проверяем разделы
                sections_count = latest_doc.get_parsed_sections_count()
                print(f"    📑 Разделов найдено: {sections_count}")
                
                # Проверяем таблицы
                tables_count = latest_doc.get_parsed_tables_count()
                print(f"    📊 Таблиц найдено: {tables_count}")
                
                if sections_count > 0 and tables_count > 0:
                    print("    ✅ Автоматическая обработка работает корректно!")
                else:
                    print("    ⚠️ Автоматическая обработка выполнена, но результаты неполные")
            else:
                print(f"    ❌ Документ не был автоматически обработан (статус: {latest_doc.status})")
        else:
            print("    ❌ Документ не найден в базе данных")
    else:
        print(f"    ❌ Ошибка загрузки: статус {response.status_code}")
        if hasattr(response, 'content'):
            print(f"    📄 Содержимое ответа: {response.content.decode()[:200]}...")
    
    # Тест 2: Проверка ручной обработки (если автоматическая не сработала)
    print("\n🔍 Тест 2: Проверка ручной обработки")
    latest_doc = DjangoDocument.objects.filter(user=user).order_by('-upload_date').first()
    if latest_doc and latest_doc.status != 'processed':
        print(f"    🔧 Запуск ручной обработки документа {latest_doc.id}...")
        
        try:
            parser_service = DocumentParserService()
            parse_result = parser_service.parse_document(latest_doc)
            
            if parse_result and 'sections' in parse_result:
                sections_count = len(parse_result['sections'])
                tables_count = len(parse_result['tables'])
                print("    ✅ Ручная обработка успешна!")
                print(f"    📑 Разделов: {sections_count}")
                print(f"    📊 Таблиц: {tables_count}")
            else:
                print("    ❌ Ошибка ручной обработки: нет данных в результате")
        except Exception as e:
            print(f"    ❌ Исключение при ручной обработке: {e}")
    else:
        print("    ⏭️ Ручная обработка не требуется")
    
    # Тест 3: Проверка интерфейса
    print("\n🔍 Тест 3: Проверка интерфейса")
    response = client.get('/')
    if response.status_code == 200:
        print("    ✅ Страница документов загружена")
        
        content = response.content.decode()
        
        # Проверяем наличие информации об автоматической обработке
        if 'Автоматическая обработка' in content:
            print("    ✅ Информация об автоматической обработке найдена")
        else:
            print("    ⚠️ Информация об автоматической обработке не найдена")
        
        # Проверяем кнопки действий
        if 'fas fa-cog' in content:
            print("    ✅ Кнопки обработки найдены")
        else:
            print("    ⚠️ Кнопки обработки не найдены")
    else:
        print(f"    ❌ Ошибка загрузки страницы: {response.status_code}")
    
    # Тест 4: Проверка страницы загрузки
    print("\n🔍 Тест 4: Проверка страницы загрузки")
    response = client.get('/upload/')
    if response.status_code == 200:
        print("    ✅ Страница загрузки доступна")
        
        content = response.content.decode()
        if 'Автоматическая обработка' in content:
            print("    ✅ Информация об автоматической обработке на странице загрузки")
        else:
            print("    ❌ Информация об автоматической обработке отсутствует на странице загрузки")
    else:
        print(f"    ❌ Страница загрузки недоступна: {response.status_code}")
    
    print("\n🎉 Тестирование автоматической обработки завершено!")
    print("=" * 60)
    
    # Итоговая статистика
    total_docs = DjangoDocument.objects.filter(user=user).count()
    processed_docs = DjangoDocument.objects.filter(user=user, status='processed').count()
    
    print(f"📊 Итоговая статистика:")
    print(f"    📄 Всего документов: {total_docs}")
    print(f"    ✅ Обработанных документов: {processed_docs}")
    print(f"    📈 Процент обработки: {(processed_docs/total_docs*100):.1f}%" if total_docs > 0 else "    📈 Процент обработки: 0%")

def main():
    """Основная функция тестирования"""
    try:
        test_auto_processing()
        return True
    except Exception as e:
        print(f"\n❌ Ошибка во время тестирования: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == '__main__':
    success = main()
    sys.exit(0 if success else 1)
