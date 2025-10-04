#!/usr/bin/env python
"""
Скрипт для создания тестового документа Word
"""
import os
import sys
import django

# Настройка Django
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'wara_project.settings')
django.setup()

from docx import Document
from docx.shared import Inches
from datetime import datetime

def create_test_document():
    """Создание тестового документа Word"""
    
    # Создаем новый документ
    doc = Document()
    
    # Заголовок документа
    title = doc.add_heading('Тестовый отчет о деятельности', 0)
    
    # Метаданные
    doc.core_properties.title = 'Тестовый отчет WARA'
    doc.core_properties.author = 'Система WARA'
    doc.core_properties.subject = 'Тестовый документ для проверки парсинга'
    doc.core_properties.comments = 'Документ создан автоматически для тестирования'
    
    # Раздел 1
    doc.add_heading('1. Введение', level=1)
    doc.add_paragraph(
        'Данный документ создан для тестирования системы автоматического анализа отчетов WARA. '
        'Он содержит различные элементы, которые должны корректно обрабатываться парсером.'
    )
    
    # Раздел 2
    doc.add_heading('2. Основные показатели', level=1)
    
    # Подраздел 2.1
    doc.add_heading('2.1 Количественные показатели', level=2)
    
    # Таблица с показателями
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Заголовки таблицы
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Показатель'
    hdr_cells[1].text = 'Текущий период'
    hdr_cells[2].text = 'Предыдущий период'
    
    # Данные таблицы
    data_rows = [
        ['Выполненные задачи', '25', '23'],
        ['Процент выполнения', '95%', '92%'],
        ['Время выполнения', '40 часов', '42 часа']
    ]
    
    for i, row_data in enumerate(data_rows, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Подраздел 2.2
    doc.add_heading('2.2 Качественные показатели', level=2)
    doc.add_paragraph(
        'Качество выполненных работ соответствует установленным стандартам. '
        'Все поставленные задачи были выполнены в срок с соблюдением требований.'
    )
    
    # Раздел 3
    doc.add_heading('3. Анализ результатов', level=1)
    
    doc.add_paragraph(
        'По результатам анализа можно сделать следующие выводы:'
    )
    
    # Список выводов
    doc.add_paragraph('1. Показатели эффективности выросли на 3%', style='List Number')
    doc.add_paragraph('2. Время выполнения задач сократилось', style='List Number')
    doc.add_paragraph('3. Качество работ остается на высоком уровне', style='List Number')
    
    # Раздел 4
    doc.add_heading('4. Планы на следующий период', level=1)
    
    doc.add_paragraph(
        'На следующий период планируется:'
    )
    
    # Маркированный список
    doc.add_paragraph('Увеличить количество выполняемых задач', style='List Bullet')
    doc.add_paragraph('Оптимизировать процессы работы', style='List Bullet')
    doc.add_paragraph('Внедрить новые технологии', style='List Bullet')
    
    # Раздел 5
    doc.add_heading('5. Заключение', level=1)
    
    doc.add_paragraph(
        'Данный тестовый документ демонстрирует основные возможности системы WARA '
        'по обработке и анализу документов Word. Парсер должен корректно извлечь:'
    )
    
    doc.add_paragraph('- Весь текстовый контент', style='List Bullet')
    doc.add_paragraph('- Структуру заголовков разных уровней', style='List Bullet')
    doc.add_paragraph('- Таблицы с данными', style='List Bullet')
    doc.add_paragraph('- Списки (нумерованные и маркированные)', style='List Bullet')
    doc.add_paragraph('- Метаданные документа', style='List Bullet')
    
    # Сохраняем документ
    filename = f'test_document_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    filepath = os.path.join('/Users/wind/Documents/WARA', filename)
    doc.save(filepath)
    
    print(f"Тестовый документ создан: {filename}")
    print(f"Путь: {filepath}")
    print(f"Размер файла: {os.path.getsize(filepath)} байт")
    
    return filepath

if __name__ == '__main__':
    create_test_document()
